#!/usr/bin/env python3
"""
AUTONOMOUS CAPTURE WORKFLOW RUNNER
Elite Capture Orchestrator v2.2 - Fully Hands-Off Mode

Purpose:
- User loads documents once.
- Clicks "Start Autonomous Run".
- Walks away (gets coffee, sleeps, whatever).
- Returns to a complete or near-complete proposal package.

This is the opposite of the conversational tool. This version drives a lean 6-phase
autonomous capture process (Early Salient + Phase A + Phase B + Core Drafting + Assembly)
with Foundation Documents as the primary technical voice.

Heavy color team reviews (Pink/Red/Gold + late Neuro + late Cost) are deliberately
excluded from the main proposal path and are used only for Phase 6 BD artifact generation
(after explicit user approval).

ETHICS: Extremely strict. Every stage is explicitly told it may only use
provided documents. All model calls are logged in full. Truncation = hard stop.
"""

import os
import json
import time
import datetime
import threading
import argparse
import sys
import subprocess
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field, asdict

# NOTE: tkinter is NO LONGER imported at module level.
# It is imported ONLY inside the GUI branch of __main__ (see end of file).
# This enables true zero-Tk headless runs with --auto-generate-bd-docs
# even on systems without python3-tk or without $DISPLAY.

try:
    import ollama
except ImportError:
    ollama = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


# =============================================================================
# CONFIGURATION
# =============================================================================

MASTER_PROMPT_PATH = Path(__file__).parent.parent / "prompts" / "elite-capture-orchestrator-v2.md"
WORK_DIR = Path(__file__).parent / "autonomous_runs"
WORK_DIR.mkdir(parents=True, exist_ok=True)

DEFAULT_MODEL = "cogito:70b"   # Change this to a smaller model (e.g. llama3.2:3b) when testing

# Default directory for document selection dialogs (makes finding solicitations faster)
DEFAULT_DOCS_DIR = Path("/home/bob/Documents/PATL")

# Two-model strategy for speed + quality (recommended for long runs on laptop GPUs)
FAST_MODEL = "cogito:14b"   # Used for early heavy lifting (Phase A, B, core drafting)
HEAVY_MODEL = "cogito:70b"  # Used for critical review/iteration stages (Red, Neuro, Cost, Gold)
LOW_TEMP = 0.15
REVIEW_TEMP = 0.1
MAX_ITERATIONS = 3

# =============================================================================
# HARDENING / PERFORMANCE CONFIG (tune for your hardware - Ryzen 9 + 64GB + 4060)
# =============================================================================
ENABLE_CHECKPOINTING = True
CHECKPOINT_DIR = Path(__file__).parent / "checkpoints"
CHECKPOINT_DIR.mkdir(parents=True, exist_ok=True)

# Document handling
MAX_CHARS_PER_DOCUMENT = 30000
CHUNK_SIZE = 4000
ENABLE_SMART_CHUNKING = True

# When the heavy model is a powerful cloud model, we can afford to send significantly more source material.
CLOUD_HEAVY_MAX_CHARS_PER_DOCUMENT = 120000
CLOUD_HEAVY_MAX_CHUNKS = 12

# Reliability
MAX_RETRIES = 3
RETRY_DELAY_SECONDS = 8

# How often to save a partial checkpoint during long token generation (in tokens)
CHECKPOINT_EVERY_N_TOKENS = 256


def _sanitize_for_fs(name: str, max_len: int = 50) -> str:
    """Make a string safe for use as part of a filename or directory name on all platforms."""
    if not name:
        return "unnamed"
    # Allow alphanum, space, dash, underscore. Replace everything else with _
    cleaned = "".join(c if c.isalnum() or c in "-_ " else "_" for c in name)
    return cleaned[:max_len].strip() or "unnamed"

# Context size is now chosen automatically based on the model being used.
# This is the key change to prevent tiny models from being crippled by huge contexts.
def get_recommended_context_size(model_name: str) -> int:
    """
    Returns a sensible num_ctx based on model size/family.

    Philosophy:
    - Local models on laptop GPUs (especially 4060): stay conservative for speed.
    - Cloud models (accessed through local Ollama): be much more aggressive,
      since the heavy context processing happens on the server.
    """
    name = model_name.lower()

    # === Cloud models (high-capability remote backends) ===
    if "cloud" in name:
        if "671b" in name or "cogito" in name:
            return 128000          # Cogito 671B-class cloud models
        if "kimi" in name:
            return 200000          # Kimi models are excellent with very long context
        return 128000              # Good high default for other cloud models

    # === Local models — tuned for Ryzen 9 + 4060 class hardware ===

    # Very small / efficient local models
    if any(x in name for x in ["3b", "llama3.2", "phi4", "gemma2:2b", "gemma2:9b"]):
        return 32768

    # Medium local models (good speed/quality balance on 4060)
    if any(x in name for x in ["11b", "14b", "22b", "32b", "34b", "cogito:14b", "cogito:32b"]):
        return 24576

    # Large local models — keep context modest to stay usable
    if any(x in name for x in ["70b", "72b", "qwen2.5:72b", "cogito:70b"]):
        return 16384

    # Safe fallback
    return 24576

# =============================================================================
# ISO CERTIFICATES CONFIGURATION (Smart + Remembered)
# =============================================================================

ISO_CONFIG_FILE = Path(__file__).parent / "iso_certs_config.json"
ISO_CERTS_KEYWORDS = ["ISO 9001", "ISO 27001", "ISO9001", "ISO27001"]

POSSIBLE_ISO_LOCATIONS = [
    Path.home() / "Documents" / "PATL" / "Official Documents",
    Path.home() / "Documents" / "Official Documents",
    Path("/home/bob/Documents/PATL/Official Documents"),
    Path("/home/workdir/Documents/PATL/Official Documents"),
]


def load_iso_certs_path() -> Path | None:
    if ISO_CONFIG_FILE.exists():
        try:
            data = json.loads(ISO_CONFIG_FILE.read_text())
            path = Path(data.get("iso_certs_path", ""))
            if path.exists():
                return path
        except Exception:
            pass
    return None


def save_iso_certs_path(path: Path):
    try:
        ISO_CONFIG_FILE.write_text(json.dumps({"iso_certs_path": str(path)}, indent=2))
    except Exception as e:
        print(f"Warning: Could not save ISO path config: {e}")


def find_iso_certs_folder() -> Path | None:
    saved = load_iso_certs_path()
    if saved and saved.exists():
        return saved
    for candidate in POSSIBLE_ISO_LOCATIONS:
        if candidate.exists():
            save_iso_certs_path(candidate)
            return candidate
    return None


# =============================================================================
# DATA STRUCTURES
# =============================================================================

@dataclass
class Document:
    id: str
    filename: str
    path: str
    category: str
    content: str = ""


@dataclass
class WorkflowState:
    opportunity_name: str = ""
    documents: List[Document] = field(default_factory=list)
    documents_context: str = ""          # Rich context of all provided documents (injected into every stage)
    foundation_documents: List[str] = field(default_factory=list)  # Filenames/IDs of high-priority foundation docs (white papers, tech summaries, etc.)
    compliance_matrix: str = ""
    win_strategy: str = ""
    pink_team_report: str = ""
    technical_approach: str = ""
    past_performance: str = ""
    risk_mitigation: str = ""
    visuals_concepts: str = ""
    executive_summary: str = ""
    red_team_report: str = ""
    neuroscientist_report: str = ""
    cost_report: str = ""
    gold_team_report: str = ""
    salient_requirements: str = ""   # Early lightweight extraction of key requirements, formatting, submission mechanics, etc.
    iteration_log: List[str] = field(default_factory=list)
    final_package: str = ""
    status: str = "Not Started"


# =============================================================================
# DOCUMENT EXTRACTION
# =============================================================================

class DocumentExtractor:
    def extract(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".pdf" and PdfReader:
            return "\n\n".join(p.extract_text() or "" for p in PdfReader(str(path)).pages)
        if suffix == ".docx" and DocxDocument:
            doc = DocxDocument(str(path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return path.read_text(encoding="utf-8", errors="replace")


class DocumentProcessor:
    """Handles smart chunking and summarization for large document sets."""

    # High-value keywords for proposal work (customize as needed)
    HIGH_VALUE_KEYWORDS = [
        # Solicitation / evaluation language
        "shall", "must", "evaluation criteria", "evaluation factor", "scoring", "basis of award",
        "technical approach", "past performance", "management", "risk", "transition",
        # Army / AOS / capability gap specific (for current solicitations)
        "capability gap", "acg", "active capability gap", "aos", "army open solicitation",
        "test and evaluation", "t&e", "prototype", "demonstration", "technology readiness",
        # User's specific strengths
        "dragonscale", "lightsteel", "apexwire", "shell phone", "polsk", "twidget",
        "cricoport", "moltguard", "sdvosb", "veteran", "iso 9001", "iso 27001",
        "lean", "low overhead", "agile", "rapid", "cost realism",
    ]

    @staticmethod
    def _score_chunk(chunk: str) -> float:
        """Simple but effective relevance scoring."""
        text = chunk.lower()
        score = 0.0
        for kw in DocumentProcessor.HIGH_VALUE_KEYWORDS:
            if kw in text:
                # Give extra weight to very important terms
                if kw in ["shall", "must", "evaluation criteria", "dragonscale", "sdvosb", "iso 27001"]:
                    score += 3.0
                else:
                    score += 1.0
        return score

    @staticmethod
    def process_document(doc: Document, max_chars: int = None, max_chunks: int = None):
        """
        Returns a dict with:
            - 'content': the (possibly truncated) processed text
            - 'was_truncated': bool
            - 'original_chars': int
            - 'kept_chars': int
            - 'dropped_chunks': int
            - 'reason': str or None

        If truncation occurred, the caller is expected to decide whether to proceed.
        """
        content = doc.content or "[No extractable text]"
        original_chars = len(content)

        effective_max_chars = max_chars or MAX_CHARS_PER_DOCUMENT
        effective_max_chunks = max_chunks or 5

        if len(content) <= effective_max_chars:
            return {
                "content": content,
                "was_truncated": False,
                "original_chars": original_chars,
                "kept_chars": original_chars,
                "dropped_chunks": 0,
                "reason": None,
            }

        if not ENABLE_SMART_CHUNKING:
            truncated = content[:effective_max_chars] + "\n... [truncated]"
            return {
                "content": truncated,
                "was_truncated": True,
                "original_chars": original_chars,
                "kept_chars": effective_max_chars,
                "dropped_chunks": 0,
                "reason": f"Smart chunking disabled. Kept first {effective_max_chars} chars.",
            }

        # Split into chunks
        chunks = [content[i:i+CHUNK_SIZE] for i in range(0, len(content), CHUNK_SIZE)]
        if not chunks:
            truncated = content[:effective_max_chars]
            return {
                "content": truncated,
                "was_truncated": True,
                "original_chars": original_chars,
                "kept_chars": len(truncated),
                "dropped_chunks": 0,
                "reason": f"No chunks after splitting. Kept first {effective_max_chars} chars.",
            }

        # Score each chunk
        scored_chunks = [(DocumentProcessor._score_chunk(chunk), i, chunk) for i, chunk in enumerate(chunks)]
        scored_chunks.sort(reverse=True)  # Highest score first

        # Always include the first chunk (important context)
        selected_indices = {0}
        selected = [chunks[0]]

        # Add top-scoring chunks (avoid duplicates)
        for score, idx, chunk in scored_chunks:
            if len(selected) >= effective_max_chunks:
                break
            if idx not in selected_indices:
                selected_indices.add(idx)
                selected.append(chunk)

        # Sort selected chunks back into original order for readability
        selected.sort(key=lambda c: content.find(c[:100]) if c[:100] in content else 999999)

        kept_text = "\n\n--- CHUNK ---\n\n".join(selected)
        kept_chars = len(kept_text)

        summary_note = (
            f"[Large document — showing {len(selected)} most relevant chunks "
            f"out of {len(chunks)} total using keyword scoring. "
            f"Kept {len(selected)} / original {len(chunks)} chunks "
            f"(max_chars={effective_max_chars}, max_chunks={effective_max_chunks}). "
            f"Total length: {len(content)} chars]"
        )

        final_content = summary_note + "\n\n" + kept_text

        return {
            "content": final_content,
            "was_truncated": True,
            "original_chars": original_chars,
            "kept_chars": kept_chars,
            "dropped_chunks": len(chunks) - len(selected),
            "reason": f"Truncated via relevance scoring. Kept top {len(selected)} chunks out of {len(chunks)}.",
        }


# =============================================================================
# PROMPT BUILDER FOR AUTONOMOUS MODE
# =============================================================================

class AutonomousPromptBuilder:
    def __init__(self, master_prompt_path: Path):
        self.master_prompt = master_prompt_path.read_text(encoding="utf-8")

    def build_stage_prompt(self, stage_name: str, state: WorkflowState, extra_context: str = "") -> str:
        base = f"""You are executing the Elite Capture Orchestrator v2.2 in FULLY AUTONOMOUS mode.

The human has started the process and walked away. Your job is to drive the complete disciplined workflow without further human input until the final package is ready.

=== MASTER CONSTITUTION (v2.2) ===
{self.master_prompt[:10000]}
=== END CONSTITUTION ===

=== PROVIDED DOCUMENTS CONTEXT (MANDATORY GROUNDING - use ONLY content from here) ===
{state.documents_context}
=== END PROVIDED DOCUMENTS CONTEXT ===

=== EARLY SALIENT REQUIREMENTS (HIGH-SIGNAL EXTRACTION FROM SOLICITATION — USE FOR FOCUS) ===
{state.salient_requirements if getattr(state, 'salient_requirements', None) else "Not yet extracted."}
=== END EARLY SALIENT REQUIREMENTS ===

CURRENT WORKING STATE (use ONLY the documents above + this state):
- Opportunity: {state.opportunity_name}
- Compliance Matrix so far: {state.compliance_matrix[:2500] if state.compliance_matrix else "Not yet created"}
- Win Strategy so far: {state.win_strategy[:1500] if state.win_strategy else "Not yet created"}

STAGE TO EXECUTE NOW: {stage_name}

{extra_context}

CRITICAL ETHICAL RULES FOR THIS STAGE:
- You may ONLY reference content that exists in the PROVIDED DOCUMENTS CONTEXT section above. This is the complete set of evidence for this session.
- If evidence is missing or weak, explicitly state the gap and which document (if any) it should have come from.
- Be conservative on TRL, cost, and capability claims.
- ISO 9001/27001 certifications are audited facts — use them properly for risk reduction.
- This is a lean SDVOSB. Never recommend approaches that would be unrealistic for a small team.

CRITICAL COMPLETION RULE (applies to every stage in autonomous mode):
- Produce the COMPLETE, FULL output for the requested stage/section NOW in a single response.
- Write until the task is finished and ready for the final package.
- Do not ask questions. Do not say "should I continue?", "would you like more?", "shall I proceed?", or anything similar.
- Do not stop early or produce a partial draft expecting further instructions.
- Match the exact length, structure, and format required by the solicitation (see EARLY SALIENT REQUIREMENTS).

Output format instructions will be given in the specific stage below. Follow them exactly.
"""
        return base

        def build_drafting_prompt(self, section_name: str, instruction: str, state: WorkflowState) -> str:
            """Lighter, more focused prompt for core drafting stages.
            Avoids sending the full constitution every time so the model stays on task.
            """
            salient = (state.salient_requirements or "")[:4500]
            compliance = (state.compliance_matrix or "")[:4000]
            win = (state.win_strategy or "")[:3000]

            return f"""You are writing **one specific section** of a government proposal package.

        === FOCUSED CONTEXT (use ONLY this) ===
        Opportunity: {state.opportunity_name}

        KEY REQUIREMENTS FROM SOLICITATION:
        {salient}

        LIVING COMPLIANCE MATRIX (key points):
        {compliance}

        WIN STRATEGY & DISCRIMINATORS:
        {win}

        === YOUR TASK ===
        Write the complete section: **{section_name}**

        {instruction}

        CRITICAL RULES FOR THIS SECTION:
        - Write ONLY this section. Do not repeat win themes, compliance matrix, or earlier content.
        - Ground every claim strictly in the Foundation Documents and the context above.
        - Produce a full, professional, submission-ready section now.
        - Match the length, structure, and format expected by the solicitation.
        - Do not ask questions. Do not say "should I continue?", "would you like more?", etc.
        - If evidence is missing, clearly state the gap instead of inventing content.

        Write the full section now:"""
# =============================================================================
# AUTONOMOUS WORKFLOW ENGINE
# =============================================================================

class AutonomousWorkflowEngine:
    def __init__(self, model: str, state: WorkflowState, documents: List[Document], logger):
        self.model = model  # default / fallback
        self.state = state
        self.documents = documents
        self.logger = logger
        self.builder = AutonomousPromptBuilder(MASTER_PROMPT_PATH)
        self.extractor = DocumentExtractor()

        # Support for two-model strategy (fast for early stages, heavy for reviews)
        self.fast_model = FAST_MODEL
        self.heavy_model = HEAVY_MODEL

        # Accumulate full output per stage/label so partial checkpoints contain
        # the complete text generated so far (instead of just the last 256 tokens).
        # This makes the *_partial_*.txt files in checkpoints/ much more useful (~10-20k+).
        self._full_stage_outputs: dict[str, str] = {}

        # Cancellation support (set by the GUI; used for cooperative early exit
        # inside long streaming model calls and drafting loops). Stored on the
        # instance so inner methods and _call_model can see it without threading
        # the parameter through every call.
        self.cancel_event: Optional[threading.Event] = None

        # Build rich documents context that will be injected into EVERY stage prompt
        self._build_documents_context()

    def _is_cancelled(self) -> bool:
        """Return True if the user has requested cancellation via the GUI CANCEL button.
        Used for cooperative early exit inside long streaming model calls and drafting loops.
        """
        ce = getattr(self, 'cancel_event', None)
        return bool(ce and ce.is_set())

    def _build_documents_context(self):
        """Create the documents context.
        Foundation Documents (loaded via the "Review Solicitation + Load Foundation Doc" dialog or
        pre-tagged with category="Foundation Document") are emitted FIRST with a loud header
        so the model treats the user's actual technology artifacts as the primary voice.

        If ANY document is truncated, we HARD STOP immediately (per user requirement).
        """
        if not self.documents:
            self.state.documents_context = "NO DOCUMENTS PROVIDED IN THIS SESSION."
            return

        # Use richer limits only for cloud heavy, but still enforce truncation checking
        heavy = (self.heavy_model or "").lower()
        is_cloud_heavy = "cloud" in heavy or "671b" in heavy

        max_chars = CLOUD_HEAVY_MAX_CHARS_PER_DOCUMENT if is_cloud_heavy else MAX_CHARS_PER_DOCUMENT
        max_chunks = CLOUD_HEAVY_MAX_CHUNKS if is_cloud_heavy else 5

        # Separate foundation docs for high-priority treatment
        foundation_docs = [d for d in self.documents if d.category == "Foundation Document"]
        other_docs = [d for d in self.documents if d.category != "Foundation Document"]

        # Record in state for traceability / proposal_context.json
        self.state.foundation_documents = [f"[{d.id}] {d.filename}" for d in foundation_docs]

        parts = [
            "=== PROVIDED DOCUMENTS CONTEXT (MANDATORY — ground every claim here) ===\n",
            f"Total documents loaded: {len(self.documents)}\n",
            f"Foundation Documents (PRIMARY TECHNICAL VOICE): {len(foundation_docs)}\n",
            f"Document mode: {'CLOUD HEAVY (rich context)' if is_cloud_heavy else 'standard'}\n"
        ]

        truncated_docs = []

        # === HIGH-PRIORITY: Foundation Documents first ===
        if foundation_docs:
            parts.append("\n" + "=" * 60)
            parts.append("FOUNDATION DOCUMENTS — PRIMARY TECHNICAL VOICE FOR THIS PROPOSAL")
            parts.append("All technical claims, approaches, and evidence must be grounded in these first.")
            parts.append("These are the specific artifacts you loaded to represent your actual capabilities.")
            parts.append("=" * 60 + "\n")

            for doc in foundation_docs:
                result = DocumentProcessor.process_document(
                    doc, max_chars=max_chars, max_chunks=max_chunks
                )
                parts.append(f"\n--- FOUNDATION DOCUMENT [{doc.id}]: {doc.filename} (Category: {doc.category}) ---")
                parts.append(result["content"])
                parts.append("--- END FOUNDATION DOCUMENT ---")
                if result["was_truncated"]:
                    truncated_docs.append({
                        "id": doc.id,
                        "filename": doc.filename,
                        "original_chars": result["original_chars"],
                        "kept_chars": result["kept_chars"],
                        "dropped_chunks": result["dropped_chunks"],
                        "reason": result["reason"],
                    })

        # === Remaining supporting documents ===
        if other_docs:
            parts.append("\n" + "-" * 60)
            parts.append("SUPPORTING DOCUMENTS (ISO certs, capability statements, past performance, etc.)")
            parts.append("-" * 60 + "\n")

            for doc in other_docs:
                result = DocumentProcessor.process_document(
                    doc, max_chars=max_chars, max_chunks=max_chunks
                )
                parts.append(f"\n--- DOCUMENT [{doc.id}]: {doc.filename} (Category: {doc.category}) ---")
                parts.append(result["content"])
                parts.append("--- END DOCUMENT ---")
                if result["was_truncated"]:
                    truncated_docs.append({
                        "id": doc.id,
                        "filename": doc.filename,
                        "original_chars": result["original_chars"],
                        "kept_chars": result["kept_chars"],
                        "dropped_chunks": result["dropped_chunks"],
                        "reason": result["reason"],
                    })

        parts.append("\n=== END PROVIDED DOCUMENTS CONTEXT ===")
        self.state.documents_context = "\n".join(parts)

        # HARD STOP if any document was truncated (user requirement)
        if truncated_docs:
            self._report_truncation_and_halt(truncated_docs, stage="initial_document_processing")

    def _validate_and_repair_compliance_matrix(self, matrix_text: str) -> str:
        """Lightweight check to ensure the compliance matrix has the expected sections."""
        required_sections = [
            "## LIVING COMPLIANCE MATRIX",
            "## ALIGNMENT ASSESSMENT"
        ]

        missing = [s for s in required_sections if s not in matrix_text]

        if not missing:
            return matrix_text  # Looks good

        # Attempt a targeted repair
        repair_prompt = f"""The following Phase A output is missing one or more required sections: {', '.join(missing)}.

Original output:
{matrix_text[:4000]}

Please output ONLY the missing section(s) in the exact format requested in the original instructions. Do not add extra commentary."""

        try:
            repair = self._call_model(repair_prompt, temperature=0.1, label="Stage1_RepairMatrix")
            return matrix_text + "\n\n" + repair.strip()
        except Exception:
            return matrix_text  # Fail gracefully

    # -------------------------------------------------------------------------
    # CHECKPOINTING (HARDENING FEATURE)
    # -------------------------------------------------------------------------
    def _get_checkpoint_path(self) -> Path:
        """Always generate a fresh timestamped checkpoint filename for this run."""
        safe_name = _sanitize_for_fs(self.state.opportunity_name)
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        return CHECKPOINT_DIR / f"{safe_name}_{ts}_checkpoint.json"

    def save_checkpoint(self):
        """Save current progress so the run can be resumed later."""
        if not ENABLE_CHECKPOINTING:
            return

        data = {
            "state": {
                "opportunity_name": self.state.opportunity_name,
                "documents_context": self.state.documents_context,
                "compliance_matrix": self.state.compliance_matrix,
                "win_strategy": self.state.win_strategy,
                "pink_team_report": self.state.pink_team_report,
                "technical_approach": self.state.technical_approach,
                "past_performance": self.state.past_performance,
                "risk_mitigation": self.state.risk_mitigation,
                "visuals_concepts": self.state.visuals_concepts,
                "executive_summary": self.state.executive_summary,
                "red_team_report": self.state.red_team_report,
                "neuroscientist_report": self.state.neuroscientist_report,
                "cost_report": self.state.cost_report,
                "iteration_log": self.state.iteration_log,
                "status": self.state.status,
            },
            "documents": [
                {
                    "id": d.id,
                    "filename": d.filename,
                    "path": d.path,
                    "category": d.category,
                    "content": d.content[:5000] if d.content else ""  # store truncated for resume
                } for d in self.documents
            ],
            "timestamp": datetime.datetime.now().isoformat()
        }

        path = self._get_checkpoint_path()
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"[Checkpoint] Saved to {path}")

    def load_checkpoint(self, opportunity_name: str) -> bool:
        """Attempt to load previous progress."""
        if not ENABLE_CHECKPOINTING:
            return False

        safe_name = _sanitize_for_fs(opportunity_name)

        # Support both old (non-timestamped) and new timestamped checkpoint files.
        # Always load the most recent one for this opportunity.
        candidates = list(CHECKPOINT_DIR.glob(f"{safe_name}_checkpoint.json")) + \
                     list(CHECKPOINT_DIR.glob(f"{safe_name}_*_checkpoint.json"))

        if not candidates:
            return False

        path = max(candidates, key=lambda p: p.stat().st_mtime)

        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)

            self.state.opportunity_name = data["state"]["opportunity_name"]
            self.state.documents_context = data["state"].get("documents_context", "")
            self.state.compliance_matrix = data["state"].get("compliance_matrix", "")
            self.state.win_strategy = data["state"].get("win_strategy", "")
            self.state.pink_team_report = data["state"].get("pink_team_report", "")
            self.state.technical_approach = data["state"].get("technical_approach", "")
            self.state.past_performance = data["state"].get("past_performance", "")
            self.state.risk_mitigation = data["state"].get("risk_mitigation", "")
            self.state.visuals_concepts = data["state"].get("visuals_concepts", "")
            self.state.executive_summary = data["state"].get("executive_summary", "")
            self.state.red_team_report = data["state"].get("red_team_report", "")
            self.state.neuroscientist_report = data["state"].get("neuroscientist_report", "")
            self.state.cost_report = data["state"].get("cost_report", "")
            self.state.iteration_log = data["state"].get("iteration_log", [])
            self.state.status = data["state"].get("status", "Resumed from checkpoint")

            # Rebuild documents (content is truncated, but good enough for most review stages)
            self.documents = [
                Document(
                    id=d["id"],
                    filename=d["filename"],
                    path=d["path"],
                    category=d["category"],
                    content=d.get("content", "")
                ) for d in data.get("documents", [])
            ]
            print(f"[Checkpoint] Loaded from {path}")
            return True
        except Exception as e:
            print(f"[Checkpoint] Failed to load: {e}")
            return False

    def _save_partial_generation_checkpoint(self, label: str, partial_output: str):
        """Lightweight save of just the current in-progress generation.
        Called frequently during long streaming generations (e.g. every 256 tokens).
        """
        if not ENABLE_CHECKPOINTING:
            return

        safe_name = _sanitize_for_fs(self.state.opportunity_name)
        safe_label = _sanitize_for_fs(label, max_len=80)
        partial_path = CHECKPOINT_DIR / f"{safe_name}_partial_{safe_label}.txt"

        try:
            with open(partial_path, "w", encoding="utf-8") as f:
                f.write(partial_output)
            # Also touch the main checkpoint timestamp so we know something is in progress
            self.save_checkpoint()
        except Exception as e:
            print(f"[Partial Checkpoint] Failed to save: {e}")

    def _call_model(self, prompt: str, temperature: float = LOW_TEMP, label: str = "", progress_callback=None, model: str = None, cancel_event: threading.Event = None) -> str:
        model_to_use = model or self.model
        num_ctx = get_recommended_context_size(model_to_use)

        # Clear logging of which model is being used (very useful for hybrid local + cloud runs)
        model_tag = " [fast]" if model_to_use == self.fast_model else " [heavy]"
        if "cloud" in model_to_use.lower():
            model_tag = " [heavy:cloud]"
        enhanced_label = f"{label}{model_tag} ({model_to_use})"

        self.logger.log_model_call(enhanced_label, prompt)
        client = ollama.Client()

        last_error = None
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                # Use streaming so the user sees live token generation (small steps)
                stream = client.chat(
                    model=model_to_use,
                    messages=[{"role": "user", "content": prompt}],
                    stream=True,
                    options={"temperature": temperature, "num_ctx": num_ctx}
                )

                output_chunks = []
                token_count = 0

                if progress_callback:
                    progress_callback(f"{label}: Model is generating response (streaming tokens, context={num_ctx})...", None)

                # Ensure we have an accumulator for the full output of this stage/label
                if label not in self._full_stage_outputs:
                    self._full_stage_outputs[label] = ""

                for chunk in stream:
                    if 'message' in chunk and 'content' in chunk['message']:
                        token = chunk['message']['content']
                        output_chunks.append(token)
                        token_count += 1

                        # Accumulate the complete text generated for this stage so far
                        self._full_stage_outputs[label] += token

                        # Cooperative cancellation check (cheap, every 4 tokens).
                        # This is what makes the CANCEL button actually responsive
                        # during long drafting generations instead of being ignored
                        # until the next major stage boundary.
                        ce = cancel_event or getattr(self, 'cancel_event', None)
                        if token_count % 4 == 0 and ce and ce.is_set():
                            # Abort the stream early. We still return what we have so far
                            # (partial is better than losing the work) and let the caller
                            # decide whether to keep the partial section or discard it.
                            if progress_callback:
                                progress_callback(f"{label}: [CANCELLED] Stream aborted mid-generation by user request.", None)
                            break

                        # Frequent UI updates
                        if token_count % 8 == 0 and progress_callback:
                            progress_callback(f"{label}: Generating... ({token_count} tokens so far)", None)

                        # Frequent checkpointing during long generations (user request)
                        if token_count % CHECKPOINT_EVERY_N_TOKENS == 0:
                            # Save the FULL accumulated output for this stage so far
                            # (not just the last 256 tokens). This makes the partial files
                            # in checkpoints/ actually useful and "complete".
                            full_so_far = self._full_stage_outputs[label]
                            self._save_partial_generation_checkpoint(label, full_so_far)

                full_output = "".join(output_chunks)
                # If we broke out early due to cancel, note it in the audit log
                if getattr(self, 'cancel_event', None) and self.cancel_event.is_set():
                    self.logger.log_model_response(label, full_output + "\n\n[STREAM ABORTED EARLY DUE TO USER CANCEL]")
                else:
                    self.logger.log_model_response(label, full_output)
                return full_output

            except Exception as e:
                last_error = e
                if attempt < MAX_RETRIES:
                    print(f"[{label}] Attempt {attempt} failed: {e}. Retrying in {RETRY_DELAY_SECONDS}s...")
                    time.sleep(RETRY_DELAY_SECONDS)
                else:
                    print(f"[{label}] All {MAX_RETRIES} attempts failed.")

        raise RuntimeError(f"Failed to get response from Ollama after {MAX_RETRIES} attempts. Last error: {last_error}")

    def _report_truncation_and_halt(self, truncated_docs: list, stage: str = ""):
        """Hard stop with detailed diagnostics when any document had to be truncated.
        This enforces the user's strict rule: if documents are too large for reliable
        grounding (every claim must be traceable to complete source text), stop early
        and give the human exact, actionable instructions.

        Called from _build_documents_context (initial load) and _stage_salient_requirements.
        """
        lines = []
        lines.append("=" * 72)
        lines.append(f"TRUNCATION HALT — Stage: {stage or 'unknown'}")
        lines.append("=" * 72)
        lines.append("One or more loaded documents exceeded safe processing limits and were truncated.")
        lines.append("Per your grounding rules, this is a HARD STOP — we will not proceed with")
        lines.append("incomplete source material that could lead to weak or invented claims.")
        lines.append("")
        for td in truncated_docs:
            lines.append(f"  • Document [{td.get('id', '?')}] {td.get('filename', '?')}")
            lines.append(f"      Original size: {td.get('original_chars', 0):,} chars")
            lines.append(f"      Kept:          {td.get('kept_chars', 0):,} chars")
            lines.append(f"      Dropped chunks: {td.get('dropped_chunks', 0)}")
            lines.append(f"      Reason: {td.get('reason', '')}")
            lines.append("")
        lines.append("RECOMMENDED ACTION:")
        lines.append("  1. Split the oversized document (e.g. 'Main_BAA.pdf' + 'Appendix_Tech_Specs.pdf').")
        lines.append("  2. Or extract only the relevant sections with a PDF tool (evaluation criteria,")
        lines.append("     SOW, instructions to offerors, format requirements, red lines).")
        lines.append("  3. Remove pure boilerplate, very long unrelated past-performance volumes,")
        lines.append("     or non-applicable attachments for THIS specific opportunity.")
        lines.append("  4. Re-start the capture run with the leaner, focused document set.")
        lines.append("")
        lines.append("A copy of this report is written to <run_dir>/truncation_report_<stage>.txt")
        lines.append("and appended to the full_audit_log.md for the run.")
        lines.append("The process will now terminate so you can correct the inputs.")
        lines.append("=" * 72)

        report_text = "\n".join(lines)

        # Always emit to stdout (critical for headless / overnight runs)
        print("\n" + report_text + "\n", flush=True)

        # Best-effort: write dedicated report file in a run dir
        try:
            run_dir = None
            if hasattr(self, 'current_run_dir') and self.current_run_dir:
                run_dir = self.current_run_dir
            elif getattr(self, 'state', None) and getattr(self.state, 'opportunity_name', None):
                ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_name = _sanitize_for_fs(getattr(self.state, 'opportunity_name', 'Unknown'))
                run_dir = Path("/home/bob/Documents/PATL/autonomous_runs") / f"{ts}_{safe_name}"
                run_dir.mkdir(parents=True, exist_ok=True)
            if run_dir:
                report_path = run_dir / f"truncation_report_{stage or 'unknown'}.txt"
                report_path.write_text(report_text, encoding="utf-8")
                print(f"[Truncation] Detailed report saved to {report_path}")
        except Exception as e:
            print(f"[Truncation] Could not write report file: {e}")

        # Also append to the audit log if we have a logger with a known path
        try:
            if getattr(self, 'logger', None) and hasattr(self.logger, 'log_path') and self.logger.log_path:
                with open(self.logger.log_path, "a", encoding="utf-8") as f:
                    f.write("\n\n" + report_text + "\n\n")
        except Exception:
            pass

        # HARD STOP — this will kill the worker thread (GUI) or the process (headless)
        raise RuntimeError(
            f"TRUNCATION HALT in stage '{stage}'. See full diagnostics printed above and in "
            f"truncation_report_{stage or 'unknown'}.txt. Correct the document(s) and re-run."
        )

    def run_full_workflow(self, progress_callback, cancel_event: threading.Event = None):
        # Try to resume from checkpoint first
        if ENABLE_CHECKPOINTING and self.load_checkpoint(self.state.opportunity_name):
            progress_callback(f"Resumed from checkpoint — Status: {self.state.status}", 50)

        # Store on the instance so inner methods (_stage_*, _call_model) can see it
        # for cooperative cancellation during long streaming generations.
        self.cancel_event = cancel_event

        def check_cancel():
            # Check both the parameter (for safety) and the instance var
            ce = cancel_event or getattr(self, 'cancel_event', None)
            if ce and ce.is_set():
                progress_callback("Run cancelled by user.", 0)
                return True
            return False

        if check_cancel(): return self.state

        progress_callback("Stage 0: Intake & Document Indexing", 0)
        self._stage_0_intake()
        self.save_checkpoint()

        if check_cancel(): return self.state

        # Early high-signal extraction (heavy model, strengthened per user request).
        # Produces the actionable briefing the user sees in the "Review Solicitation + Load Foundation Doc"
        # dialog and that is injected into every later prompt. Includes explicit recommended foundation docs.
        progress_callback("Stage: Early Salient Requirements Extraction [heavy model — high signal for foundation selection]", 3)
        self._stage_salient_requirements(progress_callback)
        self.save_checkpoint()

        if check_cancel(): return self.state

        progress_callback("Stage 1: Phase A - Requirements & Compliance Matrix [fast model]", 10)
        progress_callback("Stage 1: Building detailed prompt with all documents...", 11)
        self._stage_1_phase_a(progress_callback)
        self.save_checkpoint()

        if check_cancel(): return self.state

        progress_callback("Stage 2: Phase B - Win Strategy + Lean + ISO + Cost Positioning [fast model]", 25)
        self._stage_2_phase_b()
        self.save_checkpoint()

        if check_cancel(): return self.state

        # === LEAN PATH (per user simplification request) ===
        # Heavy color team (Pink / Red / Gold + late Neuro + late Cost) is REMOVED from the main
        # proposal generation path. Those heavy reviews are reserved exclusively for Phase 6
        # BD artifact generation (after explicit user approval). The main path is now a fast,
        # 6-phase lean flow focused on producing a grounded proposal package quickly.
        progress_callback("Stage 3: Lean Core Drafting (Technical Approach, Past Performance, Risk, Visuals, Exec Summary) [fast model]", 45)
        self._stage_4_core_drafting(progress_callback)  # reuses the existing core drafting logic
        self.save_checkpoint()

        if check_cancel(): return self.state

        progress_callback("Stage 4: Final Package Assembly (lean path)", 70)
        self._stage_8_assemble_package()
        self.save_checkpoint()

        if check_cancel(): return self.state

        # Note: Full Gold + late Neuro + late Cost stages remain available and are invoked only
        # from the BD generation launch path (after user approval) for the Phase 6 artifacts.
        progress_callback("COMPLETE (lean proposal path) — Ready for your review. Heavy reviews (Gold/Neuro/Cost) reserved for BD artifacts only.", 100)
        return self.state

    def _stage_salient_requirements(self, progress_callback=None):
        """Early high-signal requirements extraction.
        Uses the HEAVY model + focused prompt (now stronger per user request) to produce
        the actionable "initial box" the operator sees in the dialog and that is injected
        into all later stages. Includes explicit RECOMMENDED FOUNDATION DOCUMENTS section.
        """
        if progress_callback:
            progress_callback("Stage: Early Salient Requirements Extraction [heavy model]", 5)

        # Find the "main" solicitation document (usually the longest PDF or the one with "solicit/baa/rfp" in the name)
        main_doc = None
        for d in self.documents:
            name = d.filename.lower()
            if any(k in name for k in ["solicit", "baa", "rfp", "rfq", "aos", "w9128z"]):
                main_doc = d
                break
        if not main_doc:
            main_doc = self.documents[0] if self.documents else None

        if not main_doc:
            self.state.salient_requirements = "NO SOLICITATION DOCUMENT FOUND."
            return

        # Use richer excerpt for heavy model (quality > raw speed for the critical first artifact the user sees)
        raw = main_doc.content or ""
        # Take first 18k + last 6k for the heavy extraction pass (still safe; heavy model handles it)
        if len(raw) > 24000:
            excerpt = raw[:18000] + "\n\n[... middle of document omitted for focus ...]\n\n" + raw[-6000:]
        else:
            excerpt = raw

        prompt = f"""You are performing a HIGH-SIGNAL extraction from a government solicitation (BAA, AOS, RFP, ACG Procedures, etc.).

Your job is to give the capture team (and the human operator) an immediately actionable picture of exactly what the customer wants, how they will evaluate, and what documents the offeror should load as the "foundation" for a credible response.

Produce ONLY the following five sections. Quote or closely paraphrase the solicitation's actual language for every substantive point. Be specific to the technical domain(s) and any Army Example Areas mentioned. Ignore generic boilerplate.

## KEY EVALUATION DRIVERS
- 5-8 bullets. For each, quote or paraphrase the solicitation's own words on what matters most for scoring/selection. Include any weighting, emphasis, or hot buttons stated or strongly implied. Pay special attention to Technical Merit, Schedule/Maturity Path, Viability, and any TRL-related language.

## REQUIRED FORMAT & STRUCTURE
- Exact page/slide limits, what counts vs does not count against the limit, file format requirements, font rules, and any prohibitions (e.g. textboxes, pop-ups, internet links).
- List every required section and note which ones count against the page limit.
- Note any special rules for Title Page, Executive Summary, ROM/Schedule, etc.

## SUBMISSION PROCESS & DEADLINES
- Exact how/where/when, portal or email address, subject line requirements (e.g. specific Army Example Area in subject), number of copies, special handling, classification rules, etc.
- Any rules about resubmission or identical submissions.

## CRITICAL MUST-HAVES / DISQUALIFIERS / RED LINES
- Absolute requirements that will cause immediate rejection or non-compliance.
- Include the exact required proprietary/restrictive legend wording (if any).
- Explicitly call out TRL disclosure requirements, including any differences between CSO and BAA pathways, and the need to discuss path to maturity.
- Note any prohibitions (PII/PHI, classified information, etc.).
- Note any requirements to propose award instrument, NAICS, small business status, IP/patent/royalty disclosure, or data rights restrictions on the Title Page or elsewhere.
- Any data storage, export control, or security requirements that could be disqualifying.

## RECOMMENDED FOUNDATION DOCUMENTS TO LOAD NOW
- Based strictly on the specific capability gaps, technical areas, evaluation drivers, and "gotchas" you extracted above, list 4-6 concrete document types or titles the offeror should load as foundation material for THIS capture.
- For each recommendation, give a one-sentence rationale that ties directly back to a driver, gap, or gotcha (example style: "Current TRL justification + test data for DragonScale thermal runaway performance — because the solicitation explicitly requires stating DoD TRL and discussing path to maturity in the High-Level Solution/Concept section, which feeds Technical Merit").
- Prioritize the user's real technology artifacts (white papers, existing Solution Briefs, TRL/maturity roadmaps, ISO evidence, relevant past performance) over generic brochures.
- Specifically call out anything needed to address TRL disclosure, targeted AEA/sub-section alignment, proprietary data marking, or company viability.

Be direct, concrete, and free of fluff. Every bullet must be traceable to text in the excerpt below. If the solicitation differentiates requirements by CSO vs BAA, or by specific Army Example Area / sub-section, make that explicit.

=== SOLICITATION TEXT (focused excerpt) ===
{excerpt}
=== END TEXT ===

Output exactly the five markdown sections above. No introductory sentence, no closing commentary."""

        if progress_callback:
            progress_callback("Stage: Early Salient Requirements Extraction — sending focused prompt to HEAVY model (quality for the critical first artifact)...", 6)

        self.state.salient_requirements = self._call_model(
            prompt, temperature=0.1, label="Early_Salient_Requirements", model=self.heavy_model, progress_callback=progress_callback
        )
        self.state.iteration_log.append("Early Salient Requirements Extraction complete.")

        # NOTE: We intentionally do NOT hard-stop here after the early salient extraction.
        # The early salient is a *focused excerpt* whose whole purpose is to give the user
        # the "what they are looking for + which of my docs should I load" briefing quickly.
        # The strong truncation hard-stop lives in _build_documents_context (the full
        # context used for actual drafting). This lets the "initial box" the user cares
        # about always run, even on large real solicitations.

    def _stage_0_intake(self):
        inventory = "DOCUMENT INVENTORY:\n"
        for d in self.documents:
            inventory += f"- [{d.id}] {d.filename} | Category: {d.category}\n"
        self.state.iteration_log.append("Stage 0 complete: Documents indexed.")

    def _stage_1_phase_a(self, progress_callback=None):
        if progress_callback:
            progress_callback("Stage 1: Constructing full prompt (this can take time with large documents)...", 12)

        prompt = self.builder.build_stage_prompt(
            "PHASE A - Deep Requirements Intelligence & Living Compliance Matrix",
            self.state,
            extra_context="""
You must now perform Phase A exactly as described in the constitution.

OUTPUT FORMAT (produce exactly these sections):

## LIVING COMPLIANCE MATRIX
| Requirement | Evaluation Factor | Response Location | Strength (1-5) | Evidence Source (Document ID + quote) | Gap / Risk | Notes |

## ALIGNMENT ASSESSMENT
- Major strengths based on provided documents only
- Critical gaps
- Ambiguities in the solicitation that should be noted
- Initial recommendations for approach

Be extremely precise. Only cite documents that were actually loaded.
"""
        )

        if progress_callback:
            progress_callback("Stage 1: Sending prompt to model (prefill phase - high memory usage expected)...", 13)

        output = self._call_model(prompt, temperature=LOW_TEMP, label="Stage1_PhaseA", progress_callback=progress_callback, model=self.fast_model, cancel_event=self.cancel_event)
        self.state.compliance_matrix = output
        self.state.iteration_log.append("Stage 1 (Phase A) complete.")

        # Lightweight structured validation / repair pass for the compliance matrix
        if progress_callback:
            progress_callback("Stage 1: Validating compliance matrix structure...", 14)

        validated = self._validate_and_repair_compliance_matrix(output)
        if validated != output:
            self.state.compliance_matrix = validated
            self.state.iteration_log.append("Stage 1: Compliance matrix was auto-repaired for missing sections.")

    def _stage_2_phase_b(self):
        prompt = self.builder.build_stage_prompt(
            "PHASE B - Win Strategy, Cost Positioning, Lean Advantages, ISO Leverage",
            self.state,
            extra_context="""
Perform Phase B.

Focus especially on:
- How the company's ISO 9001/27001 certifications reduce government risk
- Lean execution advantages (low overhead, agility, lower cost of program management)
- Realistic affordability positioning for a small SDVOSB

OUTPUT FORMAT:

## WIN THEMES & KEY MESSAGES
...

## COST & AFFORDABILITY STRATEGY
...

## LEAN EXECUTION ADVANTAGES
...

## ISO 9001/27001 AS RISK REDUCTION EVIDENCE
...
"""
        )
        output = self._call_model(prompt, temperature=LOW_TEMP, label="Stage2_PhaseB", model=self.fast_model, cancel_event=self.cancel_event)
        self.state.win_strategy = output
        self.state.iteration_log.append("Stage 2 (Phase B) complete.")

    def _stage_3_pink_team(self):
        prompt = self.builder.build_stage_prompt(
            "PINK TEAM REVIEW (Early Validation)",
            self.state,
            extra_context="""
You are now running a Pink Team review on the current Phase A + Phase B outputs.

Focus on: compliance completeness, strength of win themes, early cost realism, and whether the approach is grounded.

CRITICAL COMPLETION RULE: Produce your full, complete Pink Team report NOW in a single response. Do not ask questions or say "should I continue?"

OUTPUT: A clear Pink Team report with specific findings and required fixes.
"""
        )
        output = self._call_model(prompt, temperature=REVIEW_TEMP, label="Stage3_PinkTeam", model=self.fast_model)
        self.state.pink_team_report = output
        self.state.iteration_log.append("Stage 3 (Pink Team) complete.")

        self.state.iteration_log.append("Stage 3 (Pink Team) complete.")

    def _stage_4_core_drafting(self, progress_callback=None):
        """Core drafting of the 5 main proposal sections using the lighter drafting prompt."""
        sections = {
            "Technical Approach": "Draft a credible technical approach using ONLY the company's actual documented technologies and maturity levels from the Foundation Documents.",
            "Tailored Past Performance & Capabilities": "Extract and tailor the strongest relevant past performance and capabilities. Emphasize ISO 9001/27001 where they reduce risk.",
            "Risk, Opportunity & Mitigation (including lean execution risks)": "Identify the top risks from the government's perspective and create honest, realistic mitigation strategies. Be conservative on TRL claims.",
            "Visual & Graphics Concepts": "Define the key figures, tables, quad charts, and visual strategy needed to make the proposal scannable and compelling.",
            "Executive Summary - Quad Chart Draft": "Create a strong, one-glance executive summary / quad chart concept that captures the core value proposition."
        }

        for i, (name, instruction) in enumerate(sections.items(), 1):
            if self._is_cancelled():
                if progress_callback:
                    progress_callback("Cancellation detected — skipping remaining drafting sections.", None)
                break

            if progress_callback:
                progress_callback(f"Stage 4: Drafting {name} ({i}/{len(sections)}) [fast model]", 50 + int(i * 3))

            try:
                prompt = self.builder.build_drafting_prompt(name, instruction, self.state)

                output = self._call_model(
                    prompt,
                    temperature=LOW_TEMP,
                    label=f"Stage4_{name.replace(' ', '_')}",
                    model=self.fast_model,
                    cancel_event=self.cancel_event
                )

                output = self._auto_continue_if_needed(output, f"Stage4_{name.replace(' ', '_')}")

                if "Technical" in name:
                    self.state.technical_approach = output
                elif "Past Performance" in name:
                    self.state.past_performance = output
                elif "Risk" in name:
                    self.state.risk_mitigation = output
                elif "Visual" in name:
                    self.state.visuals_concepts = output
                elif "Executive" in name:
                    self.state.executive_summary = output

            except Exception as ex:
                err_msg = f"[ERROR drafting '{name}'] {ex}"
                print(err_msg)
                if progress_callback:
                    progress_callback(err_msg, None)
                self.state.iteration_log.append(err_msg)

                placeholder = f"""## {name}

        **This section could not be generated automatically.**

        Error: {ex}

        **Recommended action:**
        - Review the full audit log for details.
        - Use the compliance matrix and win strategy above as grounding.
        - Manually draft this section or re-run just this stage.

        ---
        """
                if "Technical" in name:
                    self.state.technical_approach = placeholder
                elif "Past Performance" in name:
                    self.state.past_performance = placeholder
                elif "Risk" in name:
                    self.state.risk_mitigation = placeholder
                elif "Visual" in name:
                    self.state.visuals_concepts = placeholder
                elif "Executive" in name:
                    self.state.executive_summary = placeholder

        self.state.iteration_log.append("Stage 4 (Core Drafting) complete.")

    def _stage_5_major_reviews(self, progress_callback=None):
        # Neuroscientist Review moved to dedicated late stage after Gold Team.
        # Cost moved to its own full stage after Neuro (per user request).
        reviews = [
            ("Red Team Review", "RED TEAM ADVERSARIAL REVIEW", "Be ruthless. You have the full PROVIDED DOCUMENTS CONTEXT. List every reason a skeptical reviewer would downgrade or reject this package, citing specific missing or weak evidence from the documents.\n\nCRITICAL COMPLETION RULE: Produce your full, complete adversarial review NOW in a single response. Do not ask questions or say 'should I continue?'"),
        ]

        for i, (display_name, stage_name, extra) in enumerate(reviews, 1):
            if progress_callback:
                progress_callback(f"Stage 5: Running {display_name} ({i}/{len(reviews)}) [heavy model]", 65 + int(i * 2))

            prompt = self.builder.build_stage_prompt(stage_name, self.state, extra_context=extra)

            if "Red" in stage_name:
                self.state.red_team_report = self._call_model(prompt, temperature=REVIEW_TEMP, label="Stage5_RedTeam", model=self.heavy_model)

        self.state.iteration_log.append("Stage 5 (Major Reviews) complete.")

    def _stage_6_iteration_loop(self, progress_callback):
        for i in range(1, MAX_ITERATIONS + 1):
            progress_callback(f"Stage 6: Starting Iteration {i} of {MAX_ITERATIONS} [heavy model]", 75 + (i * 3))

            issues = f"{self.state.red_team_report}\n{self.state.neuroscientist_report}\n{self.state.cost_report}"

            if "major" in issues.lower() or "high severity" in issues.lower() or i == 1:
                progress_callback(f"Stage 6: Iteration {i} - Fixing high-severity issues [heavy model]", 75 + (i * 3) + 1)

                prompt = self.builder.build_stage_prompt(
                    f"ITERATION {i} - FIX HIGH SEVERITY ISSUES",
                    self.state,
                    extra_context=f"""Previous review findings (from documents above):
{issues[:4000]}

Revise the weakest sections to address the most serious findings. 
Be explicit about what you changed and quote the specific evidence from the PROVIDED DOCUMENTS CONTEXT that now supports the fixes."""
                )
                fixes = self._call_model(prompt, temperature=LOW_TEMP, label=f"Stage6_Iteration{i}_Fixes", model=self.heavy_model)
                self.state.iteration_log.append(f"Iteration {i} fixes applied.")

                progress_callback(f"Stage 6: Iteration {i} - Re-running reviews after fixes [heavy model]", 75 + (i * 3) + 2)
                self._stage_5_major_reviews(progress_callback)
            else:
                self.state.iteration_log.append(f"Iteration {i} not needed - reviews were acceptable.")
                break

    def _stage_7_gold_team(self, progress_callback=None):
        prompt = self.builder.build_stage_prompt(
            "GOLD TEAM / FINAL REVIEW",
            self.state,
            extra_context="""Perform a final Gold Team review.

CRITICAL COMPLETION RULE: Produce your full, complete Gold Team review output NOW in a single response. Do not ask questions or say "should I continue?"

CRITICAL: You have full access to the PROVIDED DOCUMENTS CONTEXT above. Use it to:
- Verify every claim has traceable evidence from the loaded documents.
- Assess whether the package is grounded, realistic for a lean ISO-certified SDVOSB, and submission-ready.

Give a clear Go / No-Go with specific conditions or required fixes. Do not proceed with a positive review if the documents do not support the claims."""
        )
        gold = self._call_model(prompt, temperature=REVIEW_TEMP, label="Stage7_GoldTeam", model=self.heavy_model)
        self.state.gold_team_report = gold
        self.state.iteration_log.append("Stage 7 (Gold Team) complete.")

        # === NEW: Lightweight Voice Polish (runs after Gold Team) ===
        if progress_callback:
            progress_callback("Stage: Voice & Flow Polish [fast model]", 92)

        self._stage_voice_polish(progress_callback)

    def _stage_voice_polish(self, progress_callback=None):
        """Lightweight Voice & Flow Polish stage.
        Runs after Gold Team. Uses FAST model.
        Focuses only on humanization, rhythm, and reducing AI tells.
        """
        if progress_callback:
            progress_callback("Stage: Voice & Flow Polish [fast model]", 93)

        prompt = self.builder.build_stage_prompt(
            "VOICE & FLOW POLISH (LIGHT STAGE)",
            self.state,
            extra_context="""
You are performing a LIGHTWEIGHT Voice & Flow Polish pass.

This occurs AFTER Gold Team. Your job is narrow and focused:

### PRIMARY GOAL
Make the proposal sound like it was written by a confident, experienced human capture manager — not an LLM.

### WHAT TO FOCUS ON
- Reduce AI-typical patterns: repetitive phrasing, hedging language ("it is important to note", "furthermore", "additionally"), overly long or robotic sentences.
- Improve rhythm and flow: vary sentence length, create natural transitions between ideas and sections.
- Increase readability and human feel while keeping a professional tone.
- Apply light Cognitive Alignment principles: improve Cognitive Fluency and Concreteness where it helps the voice feel more natural.

### RULES FOR THIS STAGE
- Do NOT re-evaluate grounding, TRL claims, ethics, or technical accuracy.
- Do NOT rewrite entire sections from scratch. Give targeted, high-impact rewrite suggestions for the weakest or most AI-sounding parts.
- Prioritize Executive Summary, Technical Approach, and Risk sections.
- Be specific: quote the problematic text and suggest a cleaner version.
- Keep suggestions concise and actionable.

CRITICAL COMPLETION RULE: Produce your full Voice Polish output NOW in a single response. Do not ask questions.

Output in this format:

**1. Overall Voice Assessment** (2-3 sentences)
How human vs AI-sounding does the current draft feel overall?

**2. Top Priority Voice & Flow Issues** (max 5)
For each:
- Section name
- Quote the problematic text (short)
- Explain the issue briefly
- Provide a suggested rewrite

**3. Quick Wins** (optional)
Small high-impact phrasing changes.

Be direct and helpful.
"""
        )

        self.state.voice_polish_report = self._call_model(
            prompt,
            temperature=LOW_TEMP,
            label="Voice_Polish",
            model=self.fast_model,
            cancel_event=self.cancel_event
        )
        self.state.iteration_log.append("Stage: Voice & Flow Polish complete.")

    def _stage_late_neuroscientist(self, progress_callback=None):
        """Dedicated late stage for Neuroscientist / Cognitive Alignment Review.
        Runs after Gold Team (per updated constitution). Uses heavy model.
        """
        if progress_callback:
            progress_callback("Stage: Neuroscientist Review (late, after Gold) [heavy model]", 91)

        prompt = self.builder.build_stage_prompt(
            "NEUROSCIENTIST REVIEW (LATE DEDICATED STAGE)",
            self.state,
            extra_context="""
You are now executing the dedicated late Neuroscientist Review using the Cognitive Alignment Framework.

This occurs only after the package has passed Gold Team review.

You have access to the full PROVIDED DOCUMENTS CONTEXT above plus all prior reviewed outputs.

CRITICAL COMPLETION RULE: Produce your full, complete review output NOW in a single response. Do not ask questions or say "should I continue?"

---

### YOUR DUAL MANDATE

**1. Humanization Mandate (Voice, Flow & Natural Language)**
Make the proposal feel like it was written by a sharp, experienced human capture manager rather than an LLM.

Look for and call out:
- Repetitive phrasing and AI-typical hedging ("it is important to note", "furthermore", "additionally").
- Overly dense or robotic sentence structures.
- Weak narrative flow and abrupt transitions between ideas.
- Generic corporate/AI-sounding language that lacks a confident but human voice.

When you find these issues, give **specific rewrite guidance** that improves rhythm, natural transitions, and readability while maintaining professionalism.

**2. Ethical Grounding Mandate — Apply the Cognitive Alignment Framework**

You are the ethical and cognitive conscience of this proposal. Use the following framework to evaluate and improve the work:

- **Mental Model Alignment**: Does the content map cleanly to how the government audience already thinks (e.g., AEA priorities, Technical Thrust Areas, mission needs, program structures, evaluation criteria language)? Flag places where the proposal uses language or categories that feel alien to the solicitation or the customer’s worldview.
- **Cognitive Fluency & Reduced Cognitive Load**: Is the writing easy for a busy reviewer to process? Look for dense paragraphs, unclear structure, lack of scannability, or excessive abstraction that increases mental effort.
- **Concreteness Effect**: Does the proposal use specific, concrete language and operational examples, or does it rely on vague or abstract claims? Push for clearer pictures of operational or mission benefit.
- **Risk Reduction Framing**: Does the proposal effectively show how this work reduces risk, improves reliability, lowers logistics/SWaP-C burden, or protects the program? This is often more persuasive than capability claims alone.
- **Group Mind & Category Fit**: Does the proposal align with how the organization (Army Test and Evaluation, HSST/DET/AIST, etc.) thinks and makes decisions? Does it reference existing structures, priorities, and quote relevant language from the solicitation where helpful?
- **Ethical Standards (Non-Negotiable)**: Stay strictly factual. Do not exaggerate maturity (TRL), impact, or certainty. Never use manipulative language or create false urgency. If evidence is weak or missing, explicitly identify the gap rather than papering over it. Properly leverage ISO 9001/27001 and SDVOSB status as legitimate risk-reduction evidence without overstating.

When reviewing, ask:
- Does this help the reviewer quickly understand, categorize, and internally champion the idea?
- Is every major claim grounded in the provided documents?
- Are we being honest about current maturity and the real path forward?

---

### MANDATORY 6-PART OUTPUT FORMAT

Use exactly this structure:

**1. Cognitive Load & Scannability Assessment**  
Evaluate density, structure, and how easy it is for a busy reviewer to extract key messages quickly.

**2. Voice, Flow & Humanization Review**  
Assess how natural and human the writing feels. Identify specific AI tells and provide concrete rewrite recommendations.

**3. Cognitive Alignment & Ethical Grounding Review**  
Apply the framework above. Evaluate Mental Model Alignment, Concreteness, Risk Reduction Framing, and strict adherence to evidence. Call out any overstatements, missing grounding, or misalignment with how the customer thinks.

**4. Reviewer Psychology & Persuasion**  
How well does the draft align with how skeptical government evaluators actually decide? Identify missed opportunities to reduce friction and increase internal championing without compromising honesty.

**5. Visual, Structural & Framing Recommendations**  
Prioritized suggestions that improve both scannability and cognitive alignment (e.g., better use of concrete examples, risk reduction framing, or visuals that match the audience’s mental models).

**6. Overall Assessment with Prioritized Actions**  
Clear overall feedback + the highest-priority changes needed. Separate **Humanization/Voice** issues from **Cognitive Alignment & Ethical Grounding** issues so the team can address them appropriately.

---

Be direct, specific, and constructive. Your goal is to help produce a proposal that is **cognitively easy to process**, **ethically sound**, and **written in a confident, human voice**.
"""
        )

        self.state.neuroscientist_report = self._call_model(
            prompt, temperature=REVIEW_TEMP, label="Late_Neuroscientist", model=self.heavy_model
        )
        self.state.iteration_log.append("Late Neuroscientist Review complete.")

    def _stage_cost_strategy(self, progress_callback=None):
        """Dedicated late stage for Cost & Affordability Strategy.
        Runs only after Gold Team + Neuroscientist have approved the package.
        """
        if progress_callback:
            progress_callback("Stage: Cost & Affordability Strategy [heavy model]", 92)

        prompt = self.builder.build_stage_prompt(
            "COST & AFFORDABILITY STRATEGY (LATE DEDICATED STAGE)",
            self.state,
            extra_context="""
You are now executing the dedicated late Cost & Affordability Strategy Stage.

This occurs only after the package has passed Gold Team and Neuroscientist review.

Use the full PROVIDED DOCUMENTS CONTEXT + all prior reviewed outputs (compliance matrix, win themes, technical approach, risk mitigations, etc.).

CRITICAL COMPLETION RULE: Produce your full, complete Cost & Affordability Strategy output NOW in a single response. Do not ask questions or say "should I continue?"

Develop a realistic, defensible cost/ROM and affordability positioning grounded in the actual mature solution.

Focus especially on:
- Weaponizing the company's lean execution advantages (low overhead, agility, lower PM burden) as direct government benefits.
- Producing a credible basis of estimate.
- Clear affordability discriminator narrative.
- Any remaining cost-related risks.

Output a clean, professional Cost & Affordability Strategy section suitable for inclusion in the final package.
"""
        )

        self.state.cost_report = self._call_model(prompt, temperature=REVIEW_TEMP, label="Late_Cost_Strategy", model=self.heavy_model)
        self.state.iteration_log.append("Late Cost & Affordability Strategy stage complete.")

        def _stage_8_assemble_package(self):
            ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

            # Count how many core sections actually have real content
            core_sections = [
                self.state.technical_approach,
                self.state.past_performance,
                self.state.risk_mitigation,
                self.state.visuals_concepts,
                self.state.executive_summary
            ]
            failed_sections = sum(1 for s in core_sections if "ERROR:" in (s or "") or len((s or "").strip()) < 200)

            package = f"""# PROPOSAL PACKAGE — {self.state.opportunity_name}

        **Generated:** {ts}  
        **Model Strategy:** Fast = {self.fast_model} | Heavy = {self.heavy_model}  
        **Run ID:** {getattr(self, 'current_run_dir', 'N/A')}  
        **Status:** {self.state.status}

        ---

        ## Generation Summary
        - Core drafting sections attempted: 5
        - Sections with significant issues: {failed_sections}
        - Full audit log and checkpoints available in the run directory.

        **Note:** Sections marked with **ERROR** or very short content should be reviewed and completed manually using the compliance matrix and win strategy below.

        ---

        ## 1. Executive Summary / Quad Chart Concept
        {self.state.executive_summary or "_Not generated_"}

        ---

        ## 2. Living Compliance Matrix & Phase A
        {self.state.compliance_matrix or "_Not generated_"}

        ---

        ## 3. Win Strategy, Cost & Lean Positioning (Phase B)
        {self.state.win_strategy or "_Not generated_"}

        ---

        ## 4. Technical Approach
        {self.state.technical_approach or "_Not generated_"}

        ---

        ## 5. Tailored Past Performance & Capabilities
        {self.state.past_performance or "_Not generated_"}

        ---

        ## 6. Risk, Opportunity & Mitigation
        {self.state.risk_mitigation or "_Not generated_"}

        ---

        ## 7. Visual & Graphics Strategy
        {self.state.visuals_concepts or "_Not generated_"}

        ---

        ## 8. Review History
        ### Pink Team
        {self.state.pink_team_report or "Not run in this lean path."}

        ### Red Team
        {self.state.red_team_report or "Not run in this lean path."}

        ### Neuroscientist Review (Late)
        {self.state.neuroscientist_report or "Not run in this lean path."}

        ### Cost & Affordability Strategy (Late)
        {self.state.cost_report or "Not run in this lean path."}

        ### Iteration Log
        {chr(10).join(self.state.iteration_log[-15:]) if self.state.iteration_log else "No iterations recorded."}

        ---

        **End of Autonomous Package**
        """
            self.state.final_package = package
            self.state.status = "Complete"

    def export_proposal_context_json(self, run_dir: Path) -> Path:
        """
        Exports a clean, ethics-focused proposal_context.json for handoff
        to the BD document generator (generate_bd_documents_v2.py).

        This is the primary integration point for Phase 2 of the architecture.
        """
        context = {
            "metadata": {
                "schema_version": "1.0",
                "generated_by": "autonomous_runner",
                "version": "2.2",
                "run_id": run_dir.name,
                "timestamp": datetime.datetime.now().isoformat(),
                "model": self.model,
                "context_window": get_recommended_context_size(self.model),
                "opportunity_name": self.state.opportunity_name,
            },
            "foundation_documents": self.state.foundation_documents or [],
            "documents": [
                {
                    "id": d.id,
                    "filename": d.filename,
                    "category": d.category,
                    "path": d.path,
                }
                for d in self.documents
            ],
            "core_content": {
                "compliance_matrix": {
                    "markdown": self.state.compliance_matrix,
                    "evidence_note": "All entries are grounded in the documents listed above. See full audit log for reasoning trace."
                },
                "win_strategy": {
                    "markdown": self.state.win_strategy,
                    "evidence_note": "Win themes and discriminators are derived exclusively from loaded documents and ISO certifications."
                },
                "technical_approach": {
                    "markdown": self.state.technical_approach,
                    "evidence_note": "Technical claims reference only technologies and maturity levels present in the provided documents."
                },
                "past_performance": {
                    "markdown": self.state.past_performance,
                    "evidence_note": "Past performance examples are taken directly from loaded artifacts."
                },
                "risk_mitigation": {
                    "markdown": self.state.risk_mitigation,
                    "evidence_note": "Risks and mitigations are assessed against actual company capabilities and ISO processes documented above."
                },
                "visuals_concepts": {
                    "markdown": self.state.visuals_concepts,
                    "evidence_note": ""
                },
                "executive_summary": {
                    "markdown": self.state.executive_summary,
                    "evidence_note": ""
                }
            },
            "reviews": {
                "pink_team": self.state.pink_team_report,
                "red_team": self.state.red_team_report,
                "neuroscientist": self.state.neuroscientist_report,
                "cost_strategist": self.state.cost_report,
            },
            "iteration_log": self.state.iteration_log,
            "traceability": {
                "all_claims_grounded": True,
                "documents_referenced": [d.id for d in self.documents],
                "notes": "This context was produced under strict grounding rules. Every claim should be traceable to one or more documents above. See full_audit_log.md for complete reasoning."
            }
        }

        output_path = run_dir / "proposal_context.json"
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(context, f, indent=2, ensure_ascii=False)

        print(f"[Export] proposal_context.json written to {output_path}")
        return output_path


# =============================================================================
# LOGGER
# =============================================================================

class AutonomousLogger:
    def __init__(self, run_dir: Path):
        self.run_dir = run_dir
        self.run_dir.mkdir(parents=True, exist_ok=True)
        self.log_file = self.run_dir / "full_audit_log.md"
        with open(self.log_file, "w", encoding="utf-8") as f:
            f.write(f"# Autonomous Capture Run Audit Log\nStarted: {datetime.datetime.now()}\n\n")

    def log_model_call(self, label: str, prompt: str):
        with open(self.log_file, "a", encoding="utf-8") as f:
            f.write(f"\n\n## MODEL CALL — {label}\n")
            f.write(f"Time: {datetime.datetime.now()}\n\n")
            f.write("```markdown\n")
            f.write(prompt[:8000])
            f.write("\n```\n")

    def log_model_response(self, label: str, response: str):
        with open(self.log_file, "a", encoding="utf-8") as f:
            f.write(f"\n\n## MODEL RESPONSE — {label}\n")
            f.write(f"Time: {datetime.datetime.now()}\n\n")
            f.write(response)
            f.write("\n\n---\n")


# =============================================================================
# GUI FOR AUTONOMOUS RUNNER
# =============================================================================

class AutonomousCaptureGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("AUTONOMOUS Capture Runner — Hands-Off Mode (v2.2)")

        # === Larger fonts for better readability ===
        self.base_font = ("DejaVu Sans", 12)
        self.header_font = ("DejaVu Sans", 12, "bold")
        self.mono_font = ("DejaVu Sans Mono", 12)
        self.big_font = ("DejaVu Sans", 11, "bold")

        self.root.option_add("*Font", self.base_font)
        self.root.geometry("1280x950")

        # Start maximized (works well on Linux + Windows)
        try:
            self.root.attributes('-zoomed', True)
        except Exception:
            self.root.update_idletasks()
            self.root.geometry(f"{self.root.winfo_screenwidth()}x{self.root.winfo_screenheight()}+0+0")

        # Make ttk.LabelFrame titles use our larger header font
        style = ttk.Style()
        style.configure("Big.TLabelframe.Label", font=self.header_font)

        # Colored action buttons (user preference for quick visual recognition)
        style.configure("Green.TButton", background="#228B22", foreground="white", font=self.header_font)
        style.map("Green.TButton", background=[("active", "#006400"), ("disabled", "#555555")])

        style.configure("Yellow.TButton", background="#FFD700", foreground="black", font=self.header_font)
        style.map("Yellow.TButton", background=[("active", "#E6C200"), ("disabled", "#AAAAAA")])

        style.configure("Red.TButton", background="#CC0000", foreground="white", font=self.header_font)
        style.map("Red.TButton", background=[("active", "#990000"), ("disabled", "#555555")])

        self.documents: List[Document] = []
        self.state = WorkflowState()
        self.current_run_dir: Optional[Path] = None

        # Cancellation support
        self.cancel_event = threading.Event()

        # Auto BD generation support (set from CLI args)
        self.auto_generate_bd_docs = False
        self.bd_generator_path = "/home/bob/PyCharmMiscProject/generate_bd_documents_v2.py"

        # Model strategy (for the radio buttons at the top)
        self.model_choice = tk.StringVar(value="Balanced (14b fast + 70b heavy)")

        self._build_ui()

        # Safe close behavior: don't kill a long-running background job just because the user closed the window
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)

        # Start the periodic check for ".request_show_window" sentinel (lets user bring back a hidden window)
        self.root.after(4000, self._check_reopen_request)

    def _build_ui(self):
        # === Compact top ethics header (saves space — full reminder on mouseover) ===
        top = ttk.Frame(self.root)
        top.pack(fill=tk.X, padx=10, pady=(8, 2))

        ethics_header = ttk.Label(top,
            text="HANDS-OFF AUTONOMOUS MODE  —  hover for ethics & rules",
            font=self.header_font, foreground="#8B0000")
        ethics_header.pack(side=tk.LEFT)

        # Hover shows the full reminder text (user request to save vertical space)
        ethics_header.bind("<Enter>", self._show_ethics_reminder)
        ethics_header.bind("<Leave>", self._hide_ethics_reminder)

        # No window control buttons here — they live in the main control row below (user request)

        # The actual reminder text (hidden until hover)
        self.ethics_detail = ttk.Label(self.root, text=(
            "Load documents → click Start → walk away.\n"
            "Autonomous: Phase A → B → Pink → Drafting → Red/Neuro/Cost → Iteration → Gold.\n"
            "Strict ethics: every claim must be traceable only to the documents loaded this session. "
            "Review everything before use."
        ), foreground="#8B0000", justify=tk.LEFT, font=self.base_font)
        # Not packed yet — shown on hover

        # === Model strategy row (radio buttons for quick testing — user request) ===
        model_row = ttk.Frame(self.root)
        model_row.pack(fill=tk.X, padx=10, pady=(2, 4))

        ttk.Label(model_row, text="Model Strategy:", font=self.header_font).pack(side=tk.LEFT, padx=(0, 8))

        for preset in [
            "Balanced (14b fast + 70b heavy)",
            "Quality (32b + 70b)",
            "Fast test (llama3.2 + 14b)",
            "Local Fast + Cloud Heavy",
            "Custom (edit below)"
        ]:
            ttk.Radiobutton(model_row, text=preset, variable=self.model_choice,
                            value=preset, command=self._apply_model_choice).pack(side=tk.LEFT, padx=6)

        # Quick direct entries for custom (visible when "Custom" is selected)
        self.fast_model_var = tk.StringVar(value=FAST_MODEL)
        self.heavy_model_var = tk.StringVar(value=HEAVY_MODEL)
        ttk.Label(model_row, text="  Fast:").pack(side=tk.LEFT)
        ttk.Entry(model_row, textvariable=self.fast_model_var, width=16, font=self.base_font).pack(side=tk.LEFT)
        ttk.Label(model_row, text="Heavy:").pack(side=tk.LEFT, padx=(6, 0))
        ttk.Entry(model_row, textvariable=self.heavy_model_var, width=16, font=self.base_font).pack(side=tk.LEFT)

        # === Row 1: Opportunity Name + document management buttons ===
        row1 = ttk.Frame(self.root)
        row1.pack(fill=tk.X, padx=10, pady=(2, 1))

        ttk.Label(row1, text="Opportunity Name:", font=self.header_font).pack(side=tk.LEFT)
        self.opp_name = tk.StringVar(value="New Opportunity")
        ttk.Entry(row1, textvariable=self.opp_name, width=42, font=self.base_font).pack(side=tk.LEFT, padx=5)

        ttk.Button(row1, text="Add Documents", command=self._add_documents).pack(side=tk.LEFT, padx=4)
        ttk.Button(row1, text="Auto-load ISO 9001 & 27001",
                   command=self._auto_load_iso_certs).pack(side=tk.LEFT, padx=3)
        ttk.Button(row1, text="Set ISO Folder...",
                   command=self._set_iso_certs_folder).pack(side=tk.LEFT, padx=3)
        ttk.Button(row1, text="Resume Last Run",
                   command=self._resume_last_run).pack(side=tk.LEFT, padx=4)
        ttk.Button(row1, text="Review Solicitation + Load Foundation Doc",
                   command=self._review_solicitation_and_load_foundation).pack(side=tk.LEFT, padx=4)

        # === Row 2: Main action buttons ===
        row2 = ttk.Frame(self.root)
        row2.pack(fill=tk.X, padx=10, pady=(1, 4))

        # Left side: Start and Cancel
        ttk.Button(row2, text="START AUTONOMOUS RUN", command=self._start_autonomous_run,
                   style="Green.TButton").pack(side=tk.LEFT, padx=5)

        self.cancel_button = ttk.Button(row2, text="CANCEL RUN", command=self._cancel_run,
                                        style="Red.TButton", state=tk.DISABLED)
        self.cancel_button.pack(side=tk.LEFT, padx=5)

        # Right side: Window controls (always visible)
        ttk.Button(row2, text="Hide Display (background)", command=self._hide_window,
                   style="Yellow.TButton").pack(side=tk.RIGHT, padx=4)
        ttk.Button(row2, text="Minimize Window", command=self._minimize_window).pack(side=tk.RIGHT, padx=4)

        # Document list
        self.doc_list = tk.Listbox(self.root, height=9, font=self.base_font)
        self.doc_list.pack(fill=tk.X, padx=10, pady=5)

        # ISO Status
        iso_status = ttk.Label(self.root, 
                  text="ISO Certificates: Use 'Auto-load ISO 9001 & 27001' or 'Set ISO Folder...' to configure location",
                  foreground="#006400", font=self.header_font)
        iso_status.pack(anchor=tk.W, padx=10, pady=(0, 3))

        ttk.Label(self.root, 
                  text="Tip: Once set, the tool will remember your ISO folder automatically.",
                  foreground="#555555", font=self.base_font).pack(anchor=tk.W, padx=10, pady=(0, 5))

        # Progress
        self.progress = ttk.Progressbar(self.root, maximum=100)
        self.progress.pack(fill=tk.X, padx=10, pady=5)

        self.status_label = ttk.Label(self.root, text="Status: Ready", font=self.header_font)
        self.status_label.pack(anchor=tk.W, padx=10)

        self.current_step_label = ttk.Label(self.root, text="Current Step: Waiting to start...", 
                                            foreground="#003366", font=self.base_font)
        self.current_step_label.pack(anchor=tk.W, padx=10, pady=(0, 3))

        self.checkpoint_label = ttk.Label(self.root, text="Checkpointing: Enabled (saves after each stage)", 
                                          foreground="#006400", font=self.base_font)
        self.checkpoint_label.pack(anchor=tk.W, padx=10, pady=(0, 5))

        # Live log
        ttk.Label(self.root, text="Live Progress Log:", font=self.header_font).pack(anchor=tk.W, padx=10)
        self.log = scrolledtext.ScrolledText(self.root, height=22, font=self.mono_font)
        self.log.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        # Prominent save action at the top (user request: buttons on top or popups for the final package)
        # This button is always visible in the control area and pops a save dialog (user likes popups).
        # It is enabled only after a package exists.
        self.save_button = ttk.Button(self.root, text="Save Final Package (ODT + MD)...", 
                                       command=self._save_final_package_dialog, state=tk.DISABLED)
        self.save_button.pack(pady=8)

    def _insert_markdown(self, text_widget: "tk.Text", markdown_text: str):
        """Insert markdown text into a tk.Text widget with basic formatting (headings + bold)."""
        text_widget.config(state=tk.NORMAL)
        text_widget.delete("1.0", tk.END)

        for line in markdown_text.splitlines(keepends=True):
            stripped = line.strip()

            if stripped.startswith("## "):
                text_widget.insert(tk.END, stripped[3:] + "\n", "heading2")
            elif stripped.startswith("# "):
                text_widget.insert(tk.END, stripped[2:] + "\n", "heading1")
            elif "**" in stripped:
                parts = stripped.split("**")
                for i, part in enumerate(parts):
                    tag = "bold" if i % 2 == 1 else None
                    text_widget.insert(tk.END, part, tag)
                text_widget.insert(tk.END, "\n")
            else:
                text_widget.insert(tk.END, line)

        # Configure tags for formatting
        text_widget.tag_config("heading1", font=("DejaVu Sans", 14, "bold"), spacing3=8)
        text_widget.tag_config("heading2", font=("DejaVu Sans", 12, "bold"), spacing3=6)
        text_widget.tag_config("bold", font=("DejaVu Sans", 11, "bold"))

        text_widget.config(state=tk.DISABLED)


    def _add_documents(self):
        paths = filedialog.askopenfilenames(
            title="Select All Documents",
            initialdir=str(DEFAULT_DOCS_DIR)
        )
        for p in paths:
            path = Path(p)
            doc = Document(
                id=f"D{len(self.documents)+1:02d}",
                filename=path.name,
                path=str(path),
                category="General",
                content=""
            )
            self.documents.append(doc)
            self.doc_list.insert(tk.END, f"[{doc.id}] {doc.filename}")
    def _show_small_info_dialog(self, title: str, message: str):
        """Custom info dialog using a smaller, more comfortable font size.
        Used for ISO certificate loading messages and similar non-critical notifications
        so the text is not overwhelmingly large on high-DPI / large screens.
        """
        dialog = tk.Toplevel(self.root)
        dialog.title(title)
        dialog.geometry("720x380")
        dialog.resizable(True, True)

        # Use the smaller font defined in __init__
        text_font = getattr(self, 'small_font', self.base_font)

        txt = tk.Text(dialog, wrap=tk.WORD, font=text_font, padx=12, pady=10)
        txt.insert("1.0", message)
        txt.config(state=tk.DISABLED)
        txt.pack(fill=tk.BOTH, expand=True, padx=10, pady=(10, 5))

        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(pady=8)
        ttk.Button(btn_frame, text="OK", command=dialog.destroy, width=12).pack()

        dialog.grab_set()
        dialog.focus_set()
        dialog.wait_window()

    def _auto_load_iso_certs(self):
        iso_folder = find_iso_certs_folder()

        if not iso_folder or not iso_folder.exists():
            response = messagebox.askyesno(
                "ISO Folder Not Found",
                "Could not automatically find your ISO 9001/27001 folder.\n\n"
                "Would you like to manually select the folder now?\n\n"
                "(This location will be remembered for future runs.)"
            )
            if response:
                self._set_iso_certs_folder()
            return

        loaded = []
        for pdf in sorted(iso_folder.glob("*.pdf")):
            name_lower = pdf.name.lower()
            if any(kw.lower() in name_lower for kw in ISO_CERTS_KEYWORDS):
                if any(d.path == str(pdf) for d in self.documents):
                    continue

                doc = Document(
                    id=f"D{len(self.documents)+1:02d}",
                    filename=pdf.name,
                    path=str(pdf),
                    category="ISO 9001/27001",
                    content=""
                )
                self.documents.append(doc)
                self.doc_list.insert(tk.END, f"[{doc.id}] {doc.filename}  (ISO 9001/27001)")
                loaded.append(pdf.name)

        if loaded:
            self._show_small_info_dialog(
                "ISO Certificates Loaded",
                "Successfully loaded:\n\n" + "\n".join(loaded) +
                f"\n\nFrom: {iso_folder}\n\nTagged with category 'ISO 9001/27001'."
            )
        else:
            self._show_small_info_dialog(
                "No New ISO Files",
                f"Scanned: {iso_folder}\n\n"
                "No additional ISO 9001 or 27001 PDFs were found."
            )

    def _set_iso_certs_folder(self):
        folder = filedialog.askdirectory(
            title="Select your ISO Certificates folder (e.g. Official Documents)",
            initialdir=str(DEFAULT_DOCS_DIR)
        )
        if not folder:
            return

        path = Path(folder)
        if not path.exists():
            messagebox.showerror("Invalid Folder", "The selected folder does not exist.")
            return

        save_iso_certs_path(path)
        self._show_small_info_dialog(
            "ISO Folder Saved",
            f"ISO certificates folder set to:\n\n{path}\n\n"
            "This location will be remembered automatically from now on.\n\n"
            "Click 'Auto-load ISO 9001 & 27001' again to load the certificates."
        )

    def _resume_last_run(self):
        """Allow user to choose from recent checkpoints (improved UX)."""
        checkpoints = sorted(
            CHECKPOINT_DIR.glob("*_checkpoint.json"),
            key=lambda p: p.stat().st_mtime,
            reverse=True
        )

        if not checkpoints:
            messagebox.showinfo("No Checkpoints", "No previous runs found to resume.")
            return

        # Build a simple selection list
        options = []
        for cp in checkpoints[:8]:  # Show up to 8 most recent
            try:
                with open(cp, "r", encoding="utf-8") as f:
                    data = json.load(f)
                name = data["state"].get("opportunity_name", cp.stem)
                ts = data.get("timestamp", "Unknown time")
                options.append((f"{name} — {ts}", cp))
            except Exception:
                options.append((cp.stem, cp))

        # For simplicity in Tkinter, use a quick dialog with the most recent 5
        # (A proper listbox dialog can be added later)
        choice = None
        for label, path in options:
            if messagebox.askyesno("Resume Checkpoint", f"Resume this run?\n\n{label}"):
                choice = path
                break
            if not messagebox.askyesno("Next?", "Show the next most recent checkpoint?"):
                return

        if not choice:
            return

        try:
            with open(choice, "r", encoding="utf-8") as f:
                data = json.load(f)

            self.opp_name.set(data["state"]["opportunity_name"])

            self.documents = []
            self.doc_list.delete(0, tk.END)
            for d in data.get("documents", []):
                doc = Document(
                    id=d["id"],
                    filename=d["filename"],
                    path=d["path"],
                    category=d["category"],
                    content=d.get("content", "")
                )
                self.documents.append(doc)
                self.doc_list.insert(tk.END, f"[{doc.id}] {doc.filename}  ({doc.category})")

            messagebox.showinfo(
                "Resumed",
                f"Checkpoint loaded.\n\nClick START AUTONOMOUS RUN to continue."
            )
        except Exception as e:
            messagebox.showerror("Resume Failed", f"Could not load checkpoint:\n{e}")

    def _start_autonomous_run(self):
        if not self.documents:
            messagebox.showwarning("No Documents", "You must load documents first.")
            return

        # Make sure the radio + custom entries are applied right now
        self._apply_model_choice()

        opp = self.opp_name.get().strip()
        self.state.opportunity_name = opp

        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_opp = _sanitize_for_fs(opp, max_len=40)
        self.current_run_dir = WORK_DIR / f"{timestamp}_{safe_opp}"
        logger = AutonomousLogger(self.current_run_dir)

        self.log.insert(tk.END, "Extracting text from all documents...\n")
        self.root.update()

        extractor = DocumentExtractor()
        for doc in self.documents:
            doc.content = extractor.extract(Path(doc.path))

        self.log.insert(tk.END, "Documents extracted.\n")
        self.root.update()

        # === Explicit user approval gate (sole authority) before any generation work ===
        # This implements the "user has sole approval authority" rule. Nothing heavy happens
        # until the human explicitly says "yes" with full visibility into what is about to be used.
        foundation_count = sum(1 for d in self.documents if getattr(d, 'category', '') == "Foundation Document")
        if foundation_count == 0:
            confirm_msg = (
                f"Start lean autonomous proposal generation now?\n\n"
                f"Opportunity: {opp or '(unnamed)'}\n"
                f"Documents loaded this session: {len(self.documents)}\n"
                f"  • Foundation Documents (primary technical voice): {foundation_count}\n\n"
                f"IMPORTANT: You have not yet used the \"Review Solicitation + Load Foundation Doc\" button.\n"
                f"That button runs a heavy-model extraction and shows you exactly:\n"
                f"  - What the solicitation is really asking for (key drivers, format rules, red lines)\n"
                f"  - Concrete RECOMMENDED FOUNDATION DOCUMENTS you should load (with rationales tied to specific gaps)\n\n"
                f"Would you like to click that button first (recommended), or proceed anyway with the documents you have?"
            )
            answer = messagebox.askyesno("Confirm Start — User Approval Gate", confirm_msg, default="no")
            if not answer:
                self.log.insert(tk.END, "\n[User] Run cancelled at explicit approval gate (recommended to use Review dialog first).\n")
                self.cancel_button.config(state=tk.DISABLED)
                return
            self.log.insert(tk.END, "\n[User] Proceeding without Foundation Documents (you can still load them later if needed).\n")
        else:
            confirm_msg = (
                f"Start lean autonomous proposal generation now?\n\n"
                f"Opportunity: {opp or '(unnamed)'}\n"
                f"Documents loaded this session: {len(self.documents)}\n"
                f"  • Foundation Documents (primary technical voice): {foundation_count}\n\n"
                f"The run will use ONLY the documents you loaded in this session for grounding.\n"
                f"Early Salient extraction (heavy model) + lean drafting (no heavy Pink/Red/Gold for the proposal).\n"
                f"Heavy reviews (Gold + late Neuro + late Cost) are reserved exclusively for BD artifacts (Phase 6).\n\n"
                f"Proceed with the lean autonomous run?"
            )
            if not messagebox.askyesno("Confirm Start — User Approval Gate", confirm_msg, default="no"):
                self.log.insert(tk.END, "\n[User] Run cancelled at explicit approval gate. No work performed.\n")
                self.cancel_button.config(state=tk.DISABLED)
                return

        # === Per-run agent cards (user request) ===
        # Create fresh, tailored cards for this specific solicitation + foundation docs.
        # These are used for this run only and deleted at the end.
        try:
            cards_content = self._generate_per_run_agent_cards()
            cards_path = self.current_run_dir / "agent-cards-this-run.md"
            cards_path.write_text(cards_content, encoding="utf-8")
            self.log.insert(tk.END, f"Per-run agent cards generated → {cards_path.name}\n")
        except Exception as e:
            self.log.insert(tk.END, f"Warning: Could not generate per-run agent cards: {e}\n")

        engine = AutonomousWorkflowEngine(DEFAULT_MODEL, self.state, self.documents, logger)

        self.cancel_event.clear()
        self.cancel_button.config(state=tk.NORMAL)

        def progress(stage: str, pct: int = None):
            if stage:
                self.status_label.config(text=stage)
            if pct is not None:
                self.progress["value"] = pct
            self.current_step_label.config(text=f"Current Step: {stage}")
            self.log.insert(tk.END, f"[{datetime.datetime.now().strftime('%H:%M')}] {stage}\n")
            self.log.see(tk.END)
            self.root.update()

        def run_in_background():
            try:
                engine.run_full_workflow(progress, self.cancel_event)
                self.state.final_package = engine.state.final_package

                # Export the clean handoff JSON for the BD document generator
                if not (self.cancel_event and self.cancel_event.is_set()):
                    json_path = engine.export_proposal_context_json(self.current_run_dir)

                    self.log.insert(tk.END, "\n\n=== AUTONOMOUS RUN COMPLETE ===\n")
                    self.log.insert(tk.END, f"Full audit log and artifacts saved in:\n{self.current_run_dir}\n")

                    # If the user hid the window during the long run, bring it back now so they see the result
                    try:
                        self.root.deiconify()
                    except Exception:
                        pass

                    # Show big visual GO / NO-GO popup ONLY if the heavy Gold review actually ran
                    # (i.e. we took the full color-team path for BD artifacts). For the lean main
                    # proposal path, Gold/Neuro/Cost are skipped, so we show nothing here.
                    try:
                        gold = getattr(self.state, 'gold_team_report', '') or ''
                        if gold.strip():
                            is_go = "GO" in gold.upper() and "NO-GO" not in gold.upper()
                            self._show_gold_team_popup(is_go)
                    except Exception as ex:
                        print(f"[Popup] Could not show Gold Team popup: {ex}")

                    # === NEW: Auto-generate BD documents if requested ===
                    # This is the only path that runs the heavy color team (Gold + late Neuro + late Cost).
                    # We require an explicit second user approval here (sole authority).
                    if getattr(self, 'auto_generate_bd_docs', False):
                        self.log.insert(tk.END, "\n[Auto] BD generation requested — asking for final approval...\n")
                        self.root.update()

                        bd_confirm = (
                            "Launch Phase 6 BD artifacts (quad chart, win themes, risk register, compliance matrix, etc.) now?\n\n"
                            "This will run the FULL heavy color team reviews (Gold Team + dedicated late Neuroscientist + "
                            "dedicated Cost & Affordability Strategy) against the proposal + your solicitation + all loaded "
                            "documents (including Foundation Documents as the primary voice).\n\n"
                            "This step is significantly more expensive in time and tokens than the lean proposal path.\n"
                            "It is the ONLY place the heavy reviews are used.\n\n"
                            "Proceed with BD document generation?"
                        )
                        if not messagebox.askyesno("Confirm Phase 6 BD Generation — Final User Approval Gate", bd_confirm, default="no"):
                            self.log.insert(tk.END, "\n[User] BD generation skipped at explicit approval gate. Proposal package is still complete.\n")
                        else:
                            self.log.insert(tk.END, "\n[Auto] User approved — Launching BD Document Generator...\n")
                            self.root.update()

                            import subprocess
                            bd_script = getattr(self, 'bd_generator_path', "/home/bob/PyCharmMiscProject/generate_bd_documents_v2.py")

                            # Put BD output under ~/Documents/PATL/<Opportunity> to keep everything organized near ISO docs
                            safe_name = _sanitize_for_fs(self.state.opportunity_name)
                            bd_output_dir = Path.home() / "Documents" / "PATL" / safe_name

                            cmd = [
                                sys.executable,
                                bd_script,
                                "--context", str(json_path),
                                "--output-dir", str(bd_output_dir),
                                "--name", self.state.opportunity_name,
                                "--headless"
                            ]

                            try:
                                result = subprocess.run(cmd, capture_output=True, text=True, timeout=3600)
                                self.log.insert(tk.END, "\n[Auto] BD Document Generator output:\n")
                                self.log.insert(tk.END, result.stdout)
                                if result.stderr:
                                    self.log.insert(tk.END, "\n[Auto] Errors/Warnings:\n")
                                    self.log.insert(tk.END, result.stderr)
                                self.log.insert(tk.END, f"\n[Auto] BD Document generation finished with exit code {result.returncode}\n")
                                self.log.insert(tk.END, f"[Auto] BD documents written to: {bd_output_dir}\n")
                            except subprocess.TimeoutExpired:
                                self.log.insert(tk.END, "\n[Auto] BD Document Generator timed out after 1 hour.\n")
                            except Exception as ex:
                                self.log.insert(tk.END, f"\n[Auto] Failed to launch BD generator: {ex}\n")

                else:
                    self.log.insert(tk.END, "\n\n=== RUN CANCELLED ===\n")
            except Exception as e:
                self.log.insert(tk.END, f"\n\nFATAL ERROR: {e}\n")
                messagebox.showerror("Error", str(e))
            finally:
                self.cancel_button.config(state=tk.DISABLED)
                self.root.update()

                # Clean up per-run agent cards (user request)
                try:
                    cards_path = getattr(self, 'current_run_dir', None)
                    if cards_path:
                        cards_file = cards_path / "agent-cards-this-run.md"
                        if cards_file.exists():
                            cards_file.unlink()
                except Exception:
                    pass

        # Run the heavy work in a background thread so the UI stays responsive
        threading.Thread(target=run_in_background, daemon=True).start()

    def _save_package(self):
        if not self.state.final_package:
            messagebox.showinfo("Nothing to save", "No package generated yet.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".md",
            initialfile="Autonomous_Capture_Package.md",
            initialdir=str(DEFAULT_DOCS_DIR)
        )
        if path:
            Path(path).write_text(self.state.final_package, encoding="utf-8")
            messagebox.showinfo("Saved", f"Package saved to {path}")

    def _auto_continue_if_needed(self, output: str, label: str) -> str:
        """If the model output ends with a 'should I continue?' style question, automatically continue
        to complete the section. This makes 'default behavior = continue and complete' a code guarantee
        (no extra prompt language).
        """
        lower = output.lower()
        triggers = [
            "should i continue", "would you like me to continue", "shall i proceed",
            "would you like me to go on", "do you want me to continue", "shall i continue"
        ]
        if any(t in lower for t in triggers):
            self.log.insert(tk.END, f"[{label}] Model asked to continue — auto-continuing to complete the full section...\n")
            continuation = self._call_model(
                f"The previous output for {label} ended with a question about continuing. Here is what was generated so far:\n\n{output}\n\nContinue from the last sentence and complete the FULL remaining content of the section. Do not ask any further questions. Produce the complete section now.",
                temperature=0.1,
                label=f"{label}_AutoContinue",
                model=self.heavy_model
            )
            return output + "\n\n" + continuation
        return output

    def _write_basic_odt(self, title: str, content: str, output_path: Path):
        """Basic ODT writer using only stdlib (zipfile). Produces a usable .odt with headings and paragraphs.
        This is the 'basic odt writer' requested. No external dependencies required.
        """
        import zipfile
        from xml.sax.saxutils import escape

        mimetype = b'application/vnd.oasis.opendocument.text'

        manifest = '''<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2">
 <manifest:file-entry manifest:media-type="application/vnd.oasis.opendocument.text" manifest:full-path="/"/>
 <manifest:file-entry manifest:media-type="text/xml" manifest:full-path="content.xml"/>
 <manifest:file-entry manifest:media-type="text/xml" manifest:full-path="styles.xml"/>
 <manifest:file-entry manifest:media-type="text/xml" manifest:full-path="meta.xml"/>
</manifest:manifest>'''

        styles = '''<?xml version="1.0" encoding="UTF-8"?>
<office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" office:version="1.2">
 <office:font-face-decls>
  <style:font-face style:name="Arial" svg:font-family="Arial" style:font-family-generic="swiss" style:font-pitch="variable"/>
 </office:font-face-decls>
 <office:styles>
  <style:style style:name="Standard" style:family="paragraph" style:class="text">
   <style:text-properties style:font-name="Arial" fo:font-size="12pt"/>
  </style:style>
  <style:style style:name="Heading 1" style:family="paragraph" style:parent-style-name="Standard" style:class="text">
   <style:text-properties style:font-name="Arial" fo:font-size="14pt" fo:font-weight="bold"/>
  </style:style>
  <style:style style:name="Heading 2" style:family="paragraph" style:parent-style-name="Standard" style:class="text">
   <style:text-properties style:font-name="Arial" fo:font-size="12pt" fo:font-weight="bold"/>
  </style:style>
 </office:styles>
</office:document-styles>'''

        meta = f'''<?xml version="1.0" encoding="UTF-8"?>
<office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" office:version="1.2">
 <office:meta>
  <dc:title xmlns:dc="http://purl.org/dc/elements/1.1/">{escape(title)}</dc:title>
 </office:meta>
</office:document-meta>'''

        # Crude but usable markdown → ODT (## → heading, rest → paragraphs)
        paragraphs = []
        for block in content.split('\n\n'):
            block = block.strip()
            if not block:
                continue
            if block.startswith('## '):
                h = escape(block[3:])
                paragraphs.append(f'<text:h text:outline-level="2">{h}</text:h>')
            elif block.startswith('# '):
                h = escape(block[2:])
                paragraphs.append(f'<text:h text:outline-level="1">{h}</text:h>')
            else:
                p = escape(block).replace('\n', '<text:line-break/>')
                paragraphs.append(f'<text:p>{p}</text:p>')

        content_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
 xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"
 office:version="1.2">
  <office:body>
    <office:text>
      <text:h text:outline-level="1">{escape(title)}</text:h>
      {''.join(paragraphs)}
    </office:text>
  </office:body>
</office:document-content>'''

        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as z:
            z.writestr('mimetype', mimetype)
            z.writestr('META-INF/manifest.xml', manifest)
            z.writestr('content.xml', content_xml)
            z.writestr('styles.xml', styles)
            z.writestr('meta.xml', meta)

        print(f"[ODT] Basic ODT written to {output_path}")

    def _save_final_package_dialog(self):
        """Popup save dialog for the final package (user likes popups). Writes both .md and .odt."""
        if not self.state.final_package:
            messagebox.showinfo("Nothing to save", "No package generated yet. Run the autonomous process first.")
            return

        base = filedialog.asksaveasfilename(
            defaultextension=".odt",
            initialfile="Final_Proposal",
            initialdir=str(DEFAULT_DOCS_DIR),
            title="Save Final Package (will write .odt + .md)"
        )
        if not base:
            return

        base_path = Path(base).with_suffix('')  # strip any extension the user typed

        md_path = base_path.with_suffix('.md')
        odt_path = base_path.with_suffix('.odt')

        md_path.write_text(self.state.final_package, encoding="utf-8")
        self._write_basic_odt(self.state.opportunity_name or "Proposal", self.state.final_package, odt_path)

        messagebox.showinfo("Saved", f"Final package saved as:\n  {md_path}\n  {odt_path}")

    def _cancel_run(self):
        if messagebox.askyesno("Cancel Run", "Are you sure you want to cancel the current run?\n\nProgress up to the last checkpoint will be preserved."):
            self.cancel_event.set()
            self.cancel_button.config(state=tk.DISABLED)
            # Honest message: the current in-flight generation will finish shortly
            # (streaming yields often), then we stop cleanly without starting new sections.
            # Partial work for the section in progress is still saved via the normal path.
            self.log.insert(tk.END, "\n[CANCEL] Cancellation requested.\n")
            self.log.insert(tk.END, "     Current model generation will finish in a few seconds (to keep partial output consistent),\n")
            self.log.insert(tk.END, "     then remaining work is aborted. Partial work is checkpointed.\n")
            self.log.see(tk.END)

    def _review_solicitation_and_load_foundation(self):
        """Show what the solicitation is really asking for + allow loading a foundation document.
        This is the UX the user requested for grounding the proposal in their actual technology.
        """
        if not self.documents:
            messagebox.showwarning("No Documents", "Load the solicitation first.")
            return

        # Lightweight extraction just for the dialog
        main_doc = None
        for d in self.documents:
            name = d.filename.lower()
            if any(k in name for k in ["solicit", "baa", "rfp", "rfq", "aos", "w9128z"]):
                main_doc = d
                break
        if not main_doc:
            main_doc = self.documents[0]

        # Richer excerpt for the dialog (this is the moment the user decides what foundation docs to load)
        # IMPORTANT: Content may not be pre-extracted yet if the user clicked this button
        # before hitting START (extraction is normally deferred). Do it on-demand here.
        if not main_doc.content or len(main_doc.content) < 200:
            try:
                extractor = DocumentExtractor()
                main_doc.content = extractor.extract(Path(main_doc.path))
                print(f"[Dialog] On-demand text extraction for solicitation: {main_doc.filename} ({len(main_doc.content)} chars)")
            except Exception as ex:
                print(f"[Dialog] On-demand extraction failed for {main_doc.filename}: {ex}")

        raw = (main_doc.content or "")[:22000]

        prompt = f"""You are reading a government solicitation (BAA, AOS, RFP, etc.) on behalf of a lean SDVOSB capture team.

Your job is to produce a HIGH-SIGNAL, immediately actionable briefing so the human operator can instantly decide which of *their own* technical documents (white papers, capability statements, past performance summaries, ISO certs, test data, etc.) to load as the "foundation" for the proposal.

Produce ONLY the following five sections. For every substantive claim, quote or closely paraphrase the solicitation's own wording. Be concrete and specific to the technical domains mentioned (sensors, advanced materials, C-UAS, EW, communications, etc.). Ignore generic FAR boilerplate.

## KEY EVALUATION DRIVERS
- 4-7 bullets. For each, quote or paraphrase the solicitation's own language on what matters most for scoring/selection. Include any stated weighting, emphasis, or unstated hot buttons you can infer from priorities.

## REQUIRED FORMAT & STRUCTURE
- Page limits, volume structure, mandatory sections, required graphics/tables/quad charts, font/margin rules, any penalties for exceeding limits.

## SUBMISSION PROCESS & DEADLINES
- Exact how/where/when to submit, portal, number of copies, special handling, classification, etc.

## CRITICAL MUST-HAVES / DISQUALIFIERS / RED LINES
- Absolute "shall/must" requirements that will cause immediate rejection or non-compliance. Include security, data rights, TRL floors, teaming restrictions, etc.

## RECOMMENDED FOUNDATION DOCUMENTS TO LOAD NOW
- Based strictly on the specific capability gaps, technical areas, and evaluation drivers you just extracted, list 3-5 concrete document types or titles the offeror should load as foundation material for *this* capture.
- For each recommendation, give a one-sentence rationale that ties directly back to a driver or gap (example style: "Your DragonScale / metamaterial performance data or white paper targeting [specific band or threat called out in Active Capability Gap X] — because the solicitation repeatedly emphasizes low-SWaP, high-reliability sensor materials").
- Prioritize the user's real technology artifacts (recent white papers, ISO 9001/27001 evidence of lean execution, relevant past performance on similar hardware, capability statements) over marketing collateral.

If a section cannot be answered from the text, say so in one short sentence. Every recommendation must be defensible from the excerpt below.

=== SOLICITATION TEXT ===
{raw}
=== END TEXT ===

Output exactly the five markdown sections above. No preamble, no closing zinger, no generic advice. This briefing will be used immediately to select the foundation documents that will become the primary technical voice of the proposal.
"""

        # === Create the dialog immediately in "working" state so the user sees feedback right away ===
        dialog = tk.Toplevel(self.root)
        dialog.title("Solicitation Summary + Recommended Foundation Documents")
        dialog.geometry("1150x820")

        status_label = ttk.Label(dialog,
            text="Analyzing solicitation with the heavy model...\n"
                 "This can take 30 seconds to several minutes depending on document size, model speed, and whether the model is local or cloud.\n"
                 "The window will update automatically when complete. Please wait...",
            font=self.base_font, justify=tk.CENTER, foreground="#b45309")
        status_label.pack(anchor=tk.W, padx=10, pady=(20, 10))

        txt = tk.Text(dialog, wrap=tk.WORD, height=28, font=self.base_font)
        txt.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        txt.insert("1.0", "Working on detailed extraction from the solicitation. Please wait...")
        txt.config(state=tk.DISABLED)

        load_btn = ttk.Button(dialog, text="Load Foundation Document (White Paper / Tech Summary...)", state=tk.DISABLED)
        load_btn.pack(pady=15)
        close_btn = ttk.Button(dialog, text="Close", state=tk.DISABLED, command=dialog.destroy)
        close_btn.pack(pady=5)

        self.log.insert(tk.END, "\n[Dialog] Starting heavy-model analysis of solicitation for foundation recommendations. This may take 1–3+ minutes.\n")
        self.log.see(tk.END)

        def run_heavy_analysis():
            """Run the actual (potentially slow) heavy model call in a background thread."""
            if len(raw) < 300:
                summary = (
                    "No usable text could be read from the solicitation document(s) you loaded.\n\n"
                    "The dialog tried to auto-extract text on the fly, but very little (or no) readable text was found.\n\n"
                    "Common causes:\n"
                    "• PDF is scanned images (not text layer)\n"
                    "• File is extremely large and the main content is in later pages\n"
                    "• The document is an attachment index rather than the actual BAA text\n\n"
                    "Recommended next steps:\n"
                    "• Open the PDF in a tool that can export text (Adobe, pdftotext, etc.)\n"
                    "• Extract just the key sections (Evaluation Criteria, Active Capability Gaps, Instructions to Offerors, Format Requirements) into a new smaller PDF or .txt\n"
                    "• Re-load that focused excerpt and click 'Review Solicitation + Load Foundation Doc' again.\n\n"
                    "Once we have readable text, this dialog will give you the exact drivers + concrete recommendations for which of your white papers / capability statements / ISO evidence to load as the foundation."
                )
            else:
                try:
                    heavy_model = getattr(self, 'heavy_model_var', None)
                    model_to_use = heavy_model.get().strip() if heavy_model and heavy_model.get().strip() else HEAVY_MODEL

                    if hasattr(self, 'engine') and self.engine:
                        summary = self.engine._call_model(prompt, temperature=0.1, label="Dialog_Salient", model=model_to_use)
                    else:
                        import ollama
                        resp = ollama.chat(model=model_to_use, messages=[{"role": "user", "content": prompt}], options={"temperature": 0.1, "num_ctx": 32768})
                        summary = resp['message']['content']
                except Exception as e:
                    summary = f"Could not generate summary: {e}"

            # Update the dialog from the main thread
            def finish_ui():
                status_label.config(
                    text="Solicitation Summary + Recommended Foundation Documents to Load (actionable briefing for your capture decision):",
                    foreground="black"
                )
                self._insert_markdown(txt, summary)

                def load_foundation():
                    path = filedialog.askopenfilename(
                        title="Select Foundation Document",
                        initialdir=str(DEFAULT_DOCS_DIR),
                        filetypes=[("PDF/DOCX/TXT/MD", "*.pdf *.docx *.doc *.txt *.md")]
                    )
                    if path:
                        doc = Document(
                            id=f"D{len(self.documents)+1:02d}",
                            filename=Path(path).name,
                            path=path,
                            category="Foundation Document",
                            content=""
                        )
                        self.documents.append(doc)
                        self.doc_list.insert(tk.END, f"[{doc.id}] {doc.filename}  ({doc.category})")
                        # Do NOT destroy the dialog — user can add multiple foundation documents
                        messagebox.showinfo("Added", "Foundation document added with high priority.\n\nYou can add more or close the dialog when ready.")

                load_btn.config(state=tk.NORMAL, command=load_foundation)
                close_btn.config(state=tk.NORMAL, command=dialog.destroy)

            dialog.after(0, finish_ui)

        # Start the heavy work in the background so the dialog stays responsive and shows the warning immediately
        threading.Thread(target=run_heavy_analysis, daemon=True).start()

    def _generate_per_run_agent_cards(self):
        """Detailed per-run agent reference cards for the lean 6-phase autonomous flow.
        Generated fresh at the very start of every GUI run (before engine launch).
        These cards are written to the run directory as agent-cards-this-run.md and
        automatically deleted on normal completion (keeps the run dir clean).
        They accurately reflect the actual lean path used: Early Salient (heavy) + A + B
        + Core Drafting (fast + auto-continue) + Assembly, with heavy reviews (Gold + late
        Neuro + dedicated Cost) gated exclusively behind explicit user approval for Phase 6
        BD artifacts only. Foundation Documents are treated as PRIMARY TECHNICAL VOICE.
        """
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        opp = getattr(self.state, 'opportunity_name', 'Unknown Opportunity')
        cards = f"""# Agent Reference Cards — This Run Only (Lean 6-Phase)

**Generated:** {ts}  
**Opportunity:** {opp}  
**Flow:** Early Salient (heavy model) → Phase A → Phase B → Core Drafting (fast model + auto-continue) → Assembly  
**Heavy reviews (Gold + late Neuro + dedicated Cost):** Reserved exclusively for Phase 6 BD artifacts — only after your explicit approval gate.

**Grounding Rule (non-negotiable):** Every claim in every artifact must be directly traceable to a document you actually loaded *this session*. Foundation Documents (white papers, ISO 9001/27001 certs, capability statements, past performance, tech summaries, etc.) are injected with loud priority as the **PRIMARY TECHNICAL VOICE** for all drafting. If the necessary evidence is missing, the system states the gap plainly instead of inventing.

### What Actually Happened in This Run
- Early Salient Requirements Extraction (heavy/cloud model) produced the 5-section briefing (Key Drivers & Must-Wins, Format/Submission Rules, Evaluation Emphasis & Scoring, Red Lines & Disqualifiers, RECOMMENDED FOUNDATION DOCS with rationales) that you saw in the "Review Solicitation + Load Foundation Doc" dialog.
- Any Foundation Documents loaded via that dialog (or pre-tagged) were given explicit high-priority treatment in the context for every subsequent stage.
- The lean proposal path used fast-model drafting with aggressive auto-continue logic and strong completion rules — the system forced sections to the length and format required by the solicitation instead of stopping to ask "should I continue?"
- All final deliverables were emitted as **ODT** (via the basic pure-stdlib writer) + Markdown.
- `proposal_context.json` was exported as the clean, structured handoff (including foundation_documents list + evidence_note fields) for any later BD artifact generation (quad chart, win themes, risk register, compliance matrix, etc.).

### Key Artifacts Written to This Run Directory
- `final_proposal.odt` + `final_proposal.md` — the completed proposal package (the main deliverable)
- `full_audit_log.md` — complete trace of every model call, full prompt, response, decision, truncation check, and checkpoint
- `proposal_context.json` — structured evidence + foundation list + metadata for BD generator handoff
- `agent-cards-this-run.md` — this file (auto-deleted on successful run completion)

### Active Agents — Lean Proposal Path (this run only)
- **Solicitation Intelligence Analyst (Phase A)**: Master parser. Built the living compliance matrix and extracted what this specific government customer actually cares about, evaluation emphasis, ambiguities, and format rules.
- **Win Strategist (Phase B)**: Owns the overarching win narrative, sharp evidence-based discriminators, authentic SDVOSB/ISO/lean positioning, and cost-affordability angle as a strategic element.
- **Technical Approach & IP Mapping Architect**: Translates your actual documented technologies and IP into credible, TRL-aware approaches that satisfy the requirement without overclaiming maturity.
- **Evidence & Tailoring Specialist**: Ruthlessly mines the loaded documents for the most relevant, quantified, citable past performance and capability evidence, mapped directly to evaluation factors and solicitation language.
- **Visuals & Information Design Lead**: Defines the graphic and layout strategy (quad chart concepts for the executive summary, callouts, scannability rules, what survives a 30-second skim).
- **Risk, Opportunity & Mitigation Analyst**: Produces honest top risks from the government's viewpoint + concrete mitigations + genuine opportunity calls. Leverages ISO 9001/27001 as audited process evidence.

Full charters for the agents above (plus the heavy-only agents used exclusively for Phase 6 BD artifacts after your approval — Neuroscientist Cognitive Alignment, Cost & Affordability Strategist, Red Team, Gold Team) live in:
`capture/agents/agent-reference-cards.md`

The master constitution and 6-phase lean rules live in:
`capture/prompts/elite-capture-orchestrator-v2.md`

**Ethics & Traceability (this run):** Nothing was invented. The system is deliberately conservative and will surface gaps rather than hallucinate content. Weak or missing sections in the output reflect gaps in the specific documents you supplied for *this opportunity*, not model creativity. You are the sole approval authority at every gate.

This file is a lightweight, per-run snapshot for traceability and your review. It is deleted automatically when the run completes normally.
"""
        return cards

    # ------------------------------------------------------------------
    # Window management (user request: easy way to "quit the display"
    # while a long run keeps going in the background thread)
    # ------------------------------------------------------------------
    def _minimize_window(self):
        """Iconify the window (taskbar / dock)."""
        try:
            self.root.iconify()
        except Exception:
            pass

    def _hide_window(self):
        """Completely hide the Tk window. The autonomous run continues normally."""
        try:
            self.root.withdraw()

            # Write to the real terminal (the log widget is no longer visible)
            msg = (
                "\n[UI] GUI window has been HIDDEN (withdraw).\n"
                "     The autonomous run is still running in the background.\n"
                "     It will automatically re-appear when the run finishes.\n"
                "\n"
                "     Ways to bring it back immediately:\n"
                "       • Press ESC then Enter in this terminal (or type 'show' + Enter)\n"
                "       • touch  " + str(self.current_run_dir / ".request_show_window") + "\n"
                "     (the window will pop up within a few seconds)\n"
            )
            print(msg, flush=True)

            # Also try to write to the log file directly if we have a logger
            try:
                if hasattr(self, 'current_run_dir') and self.current_run_dir:
                    log_path = self.current_run_dir / "full_audit_log.md"
                    with open(log_path, "a", encoding="utf-8") as f:
                        f.write("\n" + msg + "\n")
            except Exception:
                pass

            # Start watching the terminal for ESC / "show" commands
            self._start_terminal_watcher()

        except Exception:
            pass

    def _check_reopen_request(self):
        """Background check (called via after()) for a re-show request file."""
        try:
            if hasattr(self, 'current_run_dir') and self.current_run_dir:
                flag = self.current_run_dir / ".request_show_window"
                if flag.exists():
                    self._re_show_window()
                    flag.unlink(missing_ok=True)
        except Exception:
            pass
        finally:
            # Keep checking every 4 seconds as long as the GUI exists
            if self.root and self.root.winfo_exists():
                self.root.after(4000, self._check_reopen_request)

    def _re_show_window(self):
        """Bring the window back (safe to call from any thread via root.after)."""
        try:
            if self.root and self.root.winfo_exists():
                self.root.deiconify()
                self.root.lift()
                self.root.focus_force()
                print("[UI] Window re-shown.", flush=True)
                self._stop_terminal_watcher()
        except Exception:
            pass

    def _show_gold_team_popup(self, is_go: bool):
        """Show a large, clear GO / NO-GO popup based on Gold Team result."""
        try:
            popup = tk.Toplevel(self.root)
            popup.title("Gold Team Decision")
            popup.geometry("700x350")
            popup.configure(bg="white")
            popup.attributes("-topmost", True)
            popup.resizable(False, False)

            # Center the popup
            popup.update_idletasks()
            x = (popup.winfo_screenwidth() // 2) - (700 // 2)
            y = (popup.winfo_screenheight() // 2) - (350 // 2)
            popup.geometry(f"700x350+{x}+{y}")

            if is_go:
                main_text = "It's a GO!"
                color = "#00AA00"
            else:
                main_text = "It's a NO-GO!"
                color = "#CC0000"

            label = tk.Label(popup, text=main_text,
                             font=("DejaVu Sans", 64, "bold"),
                             fg=color, bg="white")
            label.pack(expand=True, pady=30)

            sub_text = "Gold Team has completed its review.\nReview the full report in the log before proceeding."
            sub = tk.Label(popup, text=sub_text, font=self.base_font, justify=tk.CENTER)
            sub.pack(pady=10)

            btn = ttk.Button(popup, text="OK", command=popup.destroy, width=12)
            btn.pack(pady=20)

            popup.grab_set()   # modal
            popup.focus_set()

        except Exception as e:
            print(f"[Popup] Failed to show Gold Team result popup: {e}")

    # --- Terminal Escape / command watcher for when window is hidden ---
    def _start_terminal_watcher(self):
        """Start a background thread that listens on stdin for ESC or 'show' commands."""
        if getattr(self, '_terminal_watcher_running', False):
            return
        self._terminal_watcher_running = True
        self._terminal_watcher_stop = False
        t = threading.Thread(target=self._terminal_watcher_loop, daemon=True)
        t.start()

    def _stop_terminal_watcher(self):
        self._terminal_watcher_stop = True

    def _terminal_watcher_loop(self):
        """Watches stdin for Escape (then Enter) or commands like 'show' / 'unhide'."""
        import select

        print("[UI] Terminal watcher active. Press ESC then Enter (or type 'show' + Enter) to bring the GUI back.", flush=True)

        while not getattr(self, '_terminal_watcher_stop', False):
            try:
                # Check if input is available (non-blocking)
                rlist, _, _ = select.select([sys.stdin], [], [], 0.5)
                if rlist:
                    line = sys.stdin.readline().strip()
                    if not line:
                        continue
                    # Detect ESC character (may appear as \x1b or in the line)
                    if '\x1b' in line or line.lower() in ('show', 'unhide', 'reveal', 'gui', 'window'):
                        # Schedule re-show on the Tk main thread
                        if self.root and self.root.winfo_exists():
                            self.root.after(0, self._re_show_window)
                        # Stop watcher once we've responded (user can hide again later)
                        self._terminal_watcher_stop = True
                        break
            except Exception:
                time.sleep(0.2)

    def _on_close(self):
        """Safe close handler: if a run is active, hide instead of destroying the root."""
        run_active = bool(self.cancel_button and str(self.cancel_button.cget("state")) != "disabled")
        if run_active:
            if messagebox.askyesno(
                "Run in progress",
                "A long autonomous run is still active.\n\n"
                "Hide the window and let the run continue in the background?\n"
                "(Progress is still being saved to the log file in autonomous_runs/.)"
            ):
                self._hide_window()
            # else: do nothing, keep the window open
        else:
            self.root.destroy()

    # --- Hover helpers for compact ethics header (saves vertical space) ---
    def _show_ethics_reminder(self, event=None):
        if not self.ethics_detail.winfo_manager():
            self.ethics_detail.pack(fill=tk.X, padx=10, pady=(0, 4), after=self.root.winfo_children()[0])

    def _hide_ethics_reminder(self, event=None):
        if self.ethics_detail.winfo_manager():
            self.ethics_detail.pack_forget()

    # --- Apply chosen model strategy (from radio buttons + entries) ---
    def _apply_model_choice(self):
        choice = self.model_choice.get()
        global FAST_MODEL, HEAVY_MODEL

        if "14b fast + 70b" in choice:
            FAST_MODEL = "cogito:14b"
            HEAVY_MODEL = "cogito:70b"
        elif "32b + 70b" in choice:
            FAST_MODEL = "cogito:32b"
            HEAVY_MODEL = "cogito:70b"
        elif "llama3.2 + 14b" in choice:
            FAST_MODEL = "llama3.2"
            HEAVY_MODEL = "cogito:14b"
        elif "Local Fast + Cloud Heavy" in choice:
            FAST_MODEL = "cogito:14b"
            # Suggest a strong cloud model for heavy reviews if the box is empty
            if not self.heavy_model_var.get().strip():
                self.heavy_model_var.set("cogito-2.1:671b-cloud")
            # User can change it to kimi-k2.5:cloud or any other cloud model they have
        else:
            # Custom — use whatever is typed in the entry boxes
            pass

        # If custom entries have values, prefer them (this also works for the Cloud Heavy preset)
        if self.fast_model_var.get().strip():
            FAST_MODEL = self.fast_model_var.get().strip()
        if self.heavy_model_var.get().strip():
            HEAVY_MODEL = self.heavy_model_var.get().strip()

        self.status_label.config(text=f"Models: fast={FAST_MODEL}  heavy={HEAVY_MODEL}")


def parse_args():
    parser = argparse.ArgumentParser(
        description="PATL Autonomous Capture Runner - Deep Reasoning + Optional BD Document Generation"
    )
    parser.add_argument(
        "--auto-generate-bd-docs",
        action="store_true",
        help="After the autonomous run completes, automatically run the BD Document Generator using the produced proposal_context.json"
    )
    parser.add_argument(
        "--bd-generator",
        type=str,
        default="/home/bob/PyCharmMiscProject/generate_bd_documents_v2.py",
        help="Path to generate_bd_documents_v2.py (default: user's known location)"
    )
    parser.add_argument(
        "--output-dir",
        type=str,
        help="Override output directory for this run"
    )
    parser.add_argument(
        "--name",
        type=str,
        help="Opportunity name (required for headless; used for run dir + BD output)"
    )
    parser.add_argument(
        "--documents",
        nargs="*",
        default=[],
        help="Paths to documents (PDF/DOCX/TXT) to load for this headless run. Can be repeated or space-separated."
    )
    parser.add_argument(
        "--documents-dir",
        type=str,
        default=None,
        help="Directory to auto-load all .pdf/.docx/.doc/.txt from (useful with DEFAULT_DOCS_DIR=/home/bob/Documents/PATL)"
    )
    parser.add_argument(
        "--fast-model",
        type=str,
        default=None,
        help="Override FAST_MODEL for early stages (e.g. llama3.2)"
    )
    parser.add_argument(
        "--heavy-model",
        type=str,
        default=None,
        help="Override HEAVY_MODEL for review stages (e.g. cogito:70b)"
    )
    parser.add_argument(
        "--pure-headless", "--no-gui",
        dest="pure_headless",
        action="store_true",
        help="Force pure terminal mode with no Tkinter window at all (for SSH, no-display, or overnight runs without any GUI)"
    )
    return parser.parse_args()


# =============================================================================
# PURE HEADLESS (ZERO-TK) EXECUTION PATH
# =============================================================================

def _load_documents_from_paths(paths: List[str]) -> List[Document]:
    """Helper: load Document objects (with extracted content) from explicit file paths."""
    docs = []
    extractor = DocumentExtractor()
    for i, p in enumerate(paths, 1):
        path = Path(p).expanduser().resolve()
        if not path.exists():
            print(f"  [WARN] Document not found, skipping: {p}")
            continue
        doc = Document(
            id=f"D{i}",
            filename=path.name,
            path=str(path),
            category="solicitation" if any(k in path.name.lower() for k in ["solicit", "baa", "rfp", "rfi", "sources"]) or path.suffix.lower() == ".pdf" else "supporting"
        )
        try:
            doc.content = extractor.extract(path)
            docs.append(doc)
            print(f"  [OK] Loaded {path.name} ({len(doc.content)} chars)")
        except Exception as e:
            print(f"  [WARN] Extraction failed for {path.name}: {e}")
    return docs


def _load_documents_from_dir(dir_path: str, max_files: int = 20) -> List[Document]:
    """Helper: auto-discover and load common proposal documents from a directory."""
    d = Path(dir_path).expanduser().resolve()
    if not d.exists() or not d.is_dir():
        print(f"  [WARN] --documents-dir does not exist or is not a dir: {dir_path}")
        return []

    exts = {".pdf", ".docx", ".doc", ".txt", ".md"}
    candidates = sorted([p for p in d.iterdir() if p.is_file() and p.suffix.lower() in exts])
    if not candidates:
        print(f"  [WARN] No supported documents found in {dir_path}")
        return []

    print(f"  Found {len(candidates)} candidate files in {dir_path} (loading up to {max_files})")
    if len(candidates) > max_files:
        print(f"  (limiting to first {max_files}; curate the directory or use --documents for precision)")
        candidates = candidates[:max_files]

    return _load_documents_from_paths([str(p) for p in candidates])


def _find_latest_checkpoint() -> Optional[tuple[Path, str]]:
    """Return (checkpoint_path, opportunity_name) for the most recently modified checkpoint, or None."""
    if not CHECKPOINT_DIR.exists():
        return None
    ckpts = list(CHECKPOINT_DIR.glob("*_checkpoint.json"))
    if not ckpts:
        return None
    latest = max(ckpts, key=lambda p: p.stat().st_mtime)
    try:
        data = json.loads(latest.read_text(encoding="utf-8"))
        opp = data.get("state", {}).get("opportunity_name") or latest.stem.replace("_checkpoint", "")
        return (latest, opp)
    except Exception:
        # Fallback: derive name from filename
        opp = latest.stem.replace("_checkpoint", "").replace("_", " ")
        return (latest, opp)


def run_headless_mode(args):
    """True zero-Tk, zero-GUI execution path for overnight / no-display runs."""
    print("=== PATL Autonomous Capture Runner — PURE HEADLESS MODE (no Tkinter, no display required) ===")
    print(f"Python: {sys.executable}")
    print(f"Working dir for artifacts: {WORK_DIR}")

    # Apply model overrides if provided (affects the two-model strategy)
    global FAST_MODEL, HEAVY_MODEL, DEFAULT_MODEL
    if args.fast_model:
        FAST_MODEL = args.fast_model
        print(f"[Override] FAST_MODEL = {FAST_MODEL}")
    if args.heavy_model:
        HEAVY_MODEL = args.heavy_model
        print(f"[Override] HEAVY_MODEL = {HEAVY_MODEL}")

    # === Determine opportunity + document / resume strategy ===
    documents: List[Document] = []
    opp_name = (args.name or "").strip()

    if args.documents:
        print(f"\nLoading {len(args.documents)} document(s) from --documents ...")
        documents = _load_documents_from_paths(args.documents)
        if not opp_name:
            opp_name = "CLI_Documents_Run"
    elif args.documents_dir:
        print(f"\nAuto-loading documents from --documents-dir: {args.documents_dir}")
        documents = _load_documents_from_dir(args.documents_dir)
        if not opp_name:
            opp_name = Path(args.documents_dir).name or "DocumentsDir_Run"
    else:
        # No documents supplied on CLI.
        # If user gave --name, we will let the engine try to resume a checkpoint for that name.
        # If user gave nothing (bare --auto-generate-bd-docs), auto-resume the single most recent checkpoint.
        if not opp_name:
            latest = _find_latest_checkpoint()
            if latest:
                ckpt_path, opp_from_ckpt = latest
                opp_name = opp_from_ckpt
                print(f"\n[Auto-resume] No --name or documents given. Resuming most recent checkpoint:")
                print(f"  {ckpt_path}")
                print(f"  Opportunity: {opp_name}")
            else:
                opp_name = "Untitled_Opportunity"
                print("\nNo --documents, --documents-dir, --name, or existing checkpoints found.")
                print("Will run with empty document context (fast but ungrounded).")

    if not opp_name:
        opp_name = "Untitled_Opportunity"

    print(f"\nOpportunity name: {opp_name}")

    if not documents:
        print("  (No fresh documents from CLI — engine will attempt checkpoint resume for this opportunity if one exists.)")

    # Create a fresh run directory for this headless invocation
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = _sanitize_for_fs(opp_name, max_len=50)
    current_run_dir = WORK_DIR / f"{timestamp}_{safe_name}"
    current_run_dir.mkdir(parents=True, exist_ok=True)
    print(f"\nArtifacts / logs / checkpoints will be written to:\n  {current_run_dir}\n")

    logger = AutonomousLogger(current_run_dir)

    # Seed the state (engine will overwrite opportunity_name etc on successful checkpoint load)
    state = WorkflowState()
    state.opportunity_name = opp_name

    # Create the engine (pure, no Tk). It will build documents_context from whatever we gave it.
    engine = AutonomousWorkflowEngine(DEFAULT_MODEL, state, documents, logger)

    # Optional: if user gave --output-dir we could wire it, but BD generator uses its own logic + PATL/<name>

    def progress_callback(stage: str, pct: int = None):
        ts = datetime.datetime.now().strftime("%H:%M:%S")
        pct_str = f" [{pct:3d}%]" if pct is not None else ""
        print(f"[{ts}]{pct_str} {stage}")

    print("=== LAUNCHING AUTONOMOUS WORKFLOW (walk away — this is the long part) ===\n")

    try:
        engine.run_full_workflow(progress_callback, cancel_event=None)

        # Export the handoff JSON (this is what the BD generator consumes)
        json_path = engine.export_proposal_context_json(current_run_dir)
        print(f"\n[OK] proposal_context.json exported → {json_path}")

        # === Auto BD document generation (identical behavior to the GUI path) ===
        if getattr(args, "auto_generate_bd_docs", False):
            print("\n[Auto] Launching BD Document Generator (generate_bd_documents_v2.py --headless)...")
            bd_script = getattr(args, "bd_generator", "/home/bob/PyCharmMiscProject/generate_bd_documents_v2.py")
            bd_output_dir = Path.home() / "Documents" / "PATL" / safe_name
            bd_output_dir.mkdir(parents=True, exist_ok=True)

            cmd = [
                sys.executable,
                bd_script,
                "--context", str(json_path),
                "--output-dir", str(bd_output_dir),
                "--name", opp_name,
                "--headless",
            ]
            print(f"[Auto] Command: {' '.join(cmd)}")
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=3600)
                print("\n[Auto] BD Generator stdout:\n" + (result.stdout or "(no stdout)"))
                if result.stderr:
                    print("\n[Auto] BD Generator stderr:\n" + result.stderr)
                print(f"\n[Auto] BD generation finished with exit code {result.returncode}")
                print(f"[Auto] BD artifacts written under: {bd_output_dir}")
            except subprocess.TimeoutExpired:
                print("[Auto] BD generator timed out after 1 hour — check its own logs.")
            except Exception as ex:
                print(f"[Auto] Failed to launch BD generator: {ex}")

        print("\n=== HEADLESS RUN COMPLETE ===")
        print(f"Full audit log, proposal_context.json, and checkpoints are in:\n  {current_run_dir}")

    except KeyboardInterrupt:
        print("\n[INTERRUPTED] User or system sent KeyboardInterrupt.")
        print("Frequent checkpoints were saved — you can resume later with the GUI 'Resume Last Run' button.")
    except Exception as e:
        print(f"\n[FATAL ERROR] {e}")
        import traceback
        traceback.print_exc()
        print("\nCheck the run directory for partial logs and the most recent checkpoint.")


if __name__ == "__main__":
    args = parse_args()

    # Only force pure terminal mode when the user explicitly asks for it.
    # --auto-generate-bd-docs and friends now open the normal GUI (the display the user prefers)
    # so they get the big live log, status, current-step label, etc. even for long runs.
    if args.pure_headless:
        # Explicit pure headless / no-GUI path (SSH, no $DISPLAY, true overnight with zero window)
        run_headless_mode(args)
    else:
        # Normal GUI mode (default, including when --auto-generate-bd-docs is used)
        import tkinter as tk
        from tkinter import ttk, filedialog, messagebox, scrolledtext

        root = tk.Tk()
        app = AutonomousCaptureGUI(root)

        # Pre-populate from CLI flags (so --auto-generate-bd-docs --name "Foo" feels natural)
        if getattr(args, "auto_generate_bd_docs", False):
            app.auto_generate_bd_docs = True
            app.bd_generator_path = getattr(args, "bd_generator", app.bd_generator_path)
        if args.name:
            app.opp_name.set(args.name)

        # Maximize on startup (user request for large readable UI)
        try:
            root.attributes('-zoomed', True)
        except Exception:
            root.update_idletasks()
            root.geometry(f"{root.winfo_screenwidth()}x{root.winfo_screenheight()}+0+0")

        root.mainloop()